# -*- coding: utf-8 -*-
"""
Generate two Solution Design Document (.docx) files:
  1. GRC Platform
  2. Internal Audit Platform
Vendor: Glimmora International  |  Contact: info@glimmora.ai
"""
import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

VENDOR = "Glimmora International"
EMAIL = "info@glimmora.ai"
TODAY = datetime.date(2026, 6, 1).strftime("%d %B %Y")

BRAND = RGBColor(0x1F, 0x3A, 0x5F)
ACCENT = RGBColor(0x2E, 0x86, 0xC1)
GREY = RGBColor(0x55, 0x55, 0x55)
CODECLR = RGBColor(0x0B, 0x52, 0x45)


# ---------- helpers ----------
def set_cell_bg(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def style_base(doc):
    st = doc.styles["Normal"]
    st.font.name = "Calibri"
    st.font.size = Pt(10.5)
    st.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
    pf = st.paragraph_format
    pf.space_after = Pt(6)
    pf.line_spacing = 1.15


def add_title_page(doc, platform, subtitle):
    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(VENDOR); r.font.size = Pt(26); r.font.bold = True; r.font.color.rgb = BRAND

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(EMAIL); r.font.size = Pt(11); r.font.color.rgb = ACCENT

    for _ in range(2):
        doc.add_paragraph()

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("SOLUTION DESIGN DOCUMENT"); r.font.size = Pt(20); r.font.bold = True
    r.font.color.rgb = RGBColor(0x22, 0x22, 0x22)

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(platform); r.font.size = Pt(16); r.font.bold = True; r.font.color.rgb = ACCENT

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(subtitle); r.font.size = Pt(11); r.font.italic = True; r.font.color.rgb = GREY

    for _ in range(6):
        doc.add_paragraph()

    meta = [
        ("Document Title", f"Solution Design — {platform}"),
        ("Prepared By", VENDOR),
        ("Contact", EMAIL),
        ("Document Date", TODAY),
        ("Version", "1.0"),
        ("Status", "Draft for Client Review"),
        ("Classification", "Confidential — Technical"),
    ]
    t = doc.add_table(rows=0, cols=2); t.alignment = WD_TABLE_ALIGNMENT.CENTER; t.style = "Table Grid"
    for k, v in meta:
        row = t.add_row().cells
        row[0].width = Inches(2.2); row[1].width = Inches(3.8)
        rp = row[0].paragraphs[0].add_run(k); rp.bold = True; rp.font.color.rgb = BRAND
        set_cell_bg(row[0], "EAF1F8")
        row[1].paragraphs[0].add_run(v)
    doc.add_page_break()


def h1(doc, text, n):
    p = doc.add_paragraph(); p.space_before = Pt(10)
    r = p.add_run(f"{n}.  {text}"); r.font.size = Pt(14); r.font.bold = True; r.font.color.rgb = BRAND
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr"); bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "2"); bottom.set(qn("w:color"), "2E86C1")
    pbdr.append(bottom); pPr.append(pbdr)


def h2(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text); r.font.size = Pt(11.5); r.font.bold = True; r.font.color.rgb = ACCENT


def para(doc, text):
    doc.add_paragraph(text)


def bullets(doc, items):
    for it in items:
        doc.add_paragraph(it, style="List Bullet")


def numbered(doc, items):
    for it in items:
        doc.add_paragraph(it, style="List Number")


def code_block(doc, lines):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd"); shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), "F2F4F5")
    pPr.append(shd)
    for i, ln in enumerate(lines):
        r = p.add_run(("" if i == 0 else "\n") + ln)
        r.font.name = "Consolas"; r.font.size = Pt(9); r.font.color.rgb = CODECLR


def table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers)); t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, htxt in enumerate(headers):
        set_cell_bg(hdr[i], "1F3A5F")
        run = hdr[i].paragraphs[0].add_run(htxt); run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF); run.font.size = Pt(10)
    for r in rows:
        cells = t.add_row().cells
        for i, val in enumerate(r):
            run = cells[i].paragraphs[0].add_run(val); run.font.size = Pt(9.5)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Inches(w)
    return t


def add_footer(doc, platform):
    p = doc.sections[0].footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"{VENDOR}  |  {EMAIL}  |  Solution Design — {platform}  |  Confidential")
    r.font.size = Pt(8); r.font.color.rgb = GREY


# =====================================================================
# SHARED TECHNICAL SECTIONS
# =====================================================================
def doc_control(doc):
    h1(doc, "Document Control", "1")
    h2(doc, "1.1  Revision History")
    table(doc, ["Version", "Date", "Author", "Description"],
          [["0.1", TODAY, VENDOR, "Initial draft"],
           ["1.0", TODAY, VENDOR, "Issued for client review"]])
    h2(doc, "1.2  Audience")
    bullets(doc, [
        "Client solution architects, IT and security teams.",
        "Project sponsors and business owners.",
        "Implementation and integration teams.",
    ])
    h2(doc, "1.3  Related Documents")
    bullets(doc, [
        "Statement of Work (SOW) for the same platform.",
        "Master Services Agreement / contract.",
        "Security & Encryption reference (SECURITY.md).",
    ])


def architecture(doc, n, platform):
    h1(doc, "Solution Architecture", str(n))
    para(doc,
         f"The {platform} is delivered as a modern, cloud-native, multi-tenant web application. "
         "It follows a layered architecture separating presentation, application/API, domain "
         "services, and data persistence, with a dedicated AI services tier for assistive and "
         "translation capabilities.")
    h2(doc, f"{n}.1  Logical Architecture Layers")
    table(doc, ["Layer", "Responsibility", "Technology"],
          [
              ["Presentation", "Responsive, multilingual UI (RTL/LTR)", "Next.js (App Router), React, Tailwind, shadcn/ui (Radix)"],
              ["API / Application", "REST endpoints, request handling, validation", "Next.js Route Handlers, Zod"],
              ["AuthN / AuthZ", "Authentication and RBAC enforcement", "NextAuth v5 (JWT sessions), permission matrix"],
              ["Domain Services", "Business logic per module", "TypeScript services, Prisma ORM"],
              ["Data Persistence", "Transactional store with field encryption", "PostgreSQL via Prisma + AES-256-GCM extension"],
              ["AI Services", "Assistant + dynamic translation", "Python backend (isolated, MongoDB store)"],
              ["Scheduling", "Reminders and automated jobs", "Scheduled cron endpoints"],
          ])
    h2(doc, f"{n}.2  Architecture Diagram (logical)")
    code_block(doc, [
        "        +------------------------------------------------------+",
        "        |                   Web Browser (User)                 |",
        "        |   EN / AR (RTL) / LV  -  Role-aware navigation       |",
        "        +-------------------------+----------------------------+",
        "                                  | HTTPS / TLS 1.2+",
        "        +-------------------------v----------------------------+",
        "        |        Next.js App (App Router)  -  Vercel/DO        |",
        "        |  Presentation  |  Route Handlers (REST)  |  NextAuth |",
        "        |  RBAC guard (withAuth)  |  Zod validation           |",
        "        +------+-------------------------------+---------------+",
        "               | Prisma (encrypt/decrypt ext)  |  HTTPS",
        "        +------v---------------+        +-------v-------------+",
        "        |  PostgreSQL (multi-  |        |  Python AI Backend  |",
        "        |  tenant, encrypted   |        |  (MongoDB store)    |",
        "        |  fileData columns)   |        |  Assistant + i18n   |",
        "        +----------------------+        +---------------------+",
    ])
    h2(doc, f"{n}.3  Multi-Tenancy")
    para(doc,
         "The platform is multi-tenant: each customer account is logically isolated. All tenant-"
         "scoped records carry a customerAccountId and every query path is scoped by the "
         "authenticated user's tenant, preventing cross-tenant data access. RBAC is layered on "
         "top of tenant isolation for fine-grained authorisation.")


def tech_stack(doc, n):
    h1(doc, "Technology Stack", str(n))
    table(doc, ["Concern", "Technology", "Notes"],
          [
              ["Framework", "Next.js 16 (App Router)", "SSR/ISR, server & route handlers"],
              ["Language", "TypeScript", "End-to-end type safety"],
              ["UI", "React, Tailwind CSS, shadcn/ui (Radix)", "Accessible component library"],
              ["Auth", "NextAuth v5", "JWT sessions, role expansion on session callback"],
              ["ORM", "Prisma", "Type-safe data access, migrations"],
              ["Database", "PostgreSQL", "Production RDBMS (multi-tenant)"],
              ["Validation", "Zod + React Hook Form", "Schema-based form & API validation"],
              ["AI Backend", "Python service (isolated)", "Assistant + dynamic translation; MongoDB"],
              ["Encryption", "AES-256-GCM (field-level)", "Sensitive Bytes columns at rest"],
              ["Hosting", "DigitalOcean", "App platform + managed PostgreSQL"],
              ["Scheduling", "Cron endpoints", "Daily due-date reminders"],
          ])


def security_section(doc, n):
    h1(doc, "Security Design", str(n))
    h2(doc, f"{n}.1  Authentication & Authorization")
    bullets(doc, [
        "NextAuth v5 with JWT sessions; credentials verified server-side.",
        "Role-Based Access Control (RBAC) defined as a permission matrix of roles x resources x actions x scopes.",
        "API routes protected by withAuth / withAuthOnly wrappers enforcing resource:action.",
        "Client-side permission hooks (usePermissions, useHasRole) drive conditional UI and permission-filtered navigation.",
        "Tenant scoping applied on every data path in addition to RBAC.",
    ])
    h2(doc, f"{n}.2  Encryption at Rest & in Transit")
    bullets(doc, [
        "Field-level encryption (AES-256-GCM) on sensitive fileData columns via a Prisma client extension — transparent encrypt on write, decrypt on read.",
        "TLS 1.2+ enforced for all network paths (browser, API, AI backend).",
        "Master key held outside source control (encrypted environment secrets); 90-day key-rotation runbook.",
        "Kill switch (ENCRYPTION_ENABLED) to stage encryption rollout per environment.",
        "Secrets (DB credentials, NEXTAUTH_SECRET, API secrets) never committed; managed as encrypted env vars.",
    ])
    h2(doc, f"{n}.3  Auditability & Data Protection")
    bullets(doc, [
        "Time-stamped audit trail of create/edit/delete actions.",
        "Sensitive values redacted in logs via a safe-logging utility.",
        "Compliance alignment: ISO 27001, SOC 2, GDPR and PCI considerations documented.",
    ])


def integration_section(doc, n, extra_rows=None):
    h1(doc, "Integration & Interfaces", str(n))
    para(doc,
         "Integrations are confirmed during the Design phase. The platform exposes RESTful APIs "
         "and supports standard interface patterns:")
    rows = [
        ["Authentication", "SSO / Identity Provider", "Optional SAML/OIDC integration (Design-phase)"],
        ["AI Services", "Internal HTTPS", "Assistant & dynamic translation backend"],
        ["Notifications", "Email / messaging", "Due-date reminders and alerts"],
        ["Import / Export", "File-based (CSV/XLSX)", "Bulk data load and reporting exports"],
    ]
    if extra_rows:
        rows.extend(extra_rows)
    table(doc, ["Interface", "Type", "Description"], rows)


def data_model_intro(doc, n, platform):
    h1(doc, "Data Model", str(n))
    para(doc,
         f"The {platform} data model is implemented in Prisma and persisted to PostgreSQL. "
         "All tenant-scoped entities reference a customer account for isolation, and sensitive "
         "binary content is stored in encrypted columns. The core entities are summarised below.")


def i18n_section(doc, n):
    h1(doc, "Internationalisation (i18n)", str(n))
    bullets(doc, [
        "Three languages supported: English (default, LTR), Arabic (RTL), Latvian (LTR).",
        "Static UI strings translated via a phrase-based t() function and language context.",
        "Full RTL layout support using directional styling for Arabic.",
        "Dynamic, user-entered data translated via the Python AI backend and stored per record/field/locale.",
        "Translation triggered only on create/edit; viewing never auto-translates.",
        "Multi-language data entry: users may enter data in any supported language; the system translates to the others.",
    ])


def nfr_section(doc, n):
    h1(doc, "Non-Functional Requirements", str(n))
    table(doc, ["Category", "Target / Approach"],
          [
              ["Performance", "Server-rendered pages; paginated, indexed queries for large datasets."],
              ["Scalability", "Stateless app tier; horizontal scaling on the hosting platform."],
              ["Availability", "Managed hosting with backups; target SLA agreed in contract."],
              ["Security", "RBAC, encryption at rest/in transit, audit trail, secret management."],
              ["Maintainability", "Typed codebase, modular structure, automated migrations."],
              ["Usability", "Responsive, accessible (Radix), multilingual UI."],
              ["Compliance", "ISO 27001 / SOC 2 / GDPR / PCI alignment."],
              ["Recoverability", "Database backups and documented restore procedure."],
          ])


def deployment_section(doc, n):
    h1(doc, "Deployment & Environments", str(n))
    h2(doc, f"{n}.1  Environments")
    table(doc, ["Environment", "Purpose", "Notes"],
          [
              ["Development", "Engineering & integration", "Local / isolated database"],
              ["UAT / Staging", "Client acceptance testing", "Production-like; env-gated features"],
              ["Production", "Live operations", "Managed PostgreSQL, encrypted secrets"],
          ])
    h2(doc, f"{n}.2  Deployment Approach")
    bullets(doc, [
        "Cloud-native deployment on DigitalOcean (app platform + managed PostgreSQL).",
        "Database schema managed via Prisma migrations.",
        "Environment-specific configuration through encrypted environment variables.",
        "Feature/risk gating via environment flags (e.g. encryption kill switch) to validate on UAT before Production.",
        "Scheduled jobs (due-date reminders) run via secured cron endpoints.",
    ])


def assumptions_risks(doc, n):
    h1(doc, "Assumptions, Dependencies & Risks", str(n))
    h2(doc, f"{n}.1  Assumptions & Dependencies")
    bullets(doc, [
        "Final integration endpoints and identity provider details are confirmed in the Design phase.",
        "The Client provides structured source/reference data for migration.",
        "Hosting region, sizing and SLA targets are agreed before Production provisioning.",
        "The Client provides timely UAT and sign-off.",
    ])
    h2(doc, f"{n}.2  Key Risks & Mitigations")
    table(doc, ["Risk", "Mitigation"],
          [
              ["Scope creep", "Baseline sign-off + Change Control."],
              ["Data quality issues", "Early data profiling and cleansing window."],
              ["Integration delays", "Confirm interfaces in Design; mock early."],
              ["Adoption resistance", "Role-based training and enablement."],
          ])


# =====================================================================
# DOCUMENT 1 — GRC PLATFORM
# =====================================================================
def build_grc():
    doc = Document(); style_base(doc)
    platform = "Governance, Risk & Compliance (GRC) Platform"
    add_title_page(doc, "GRC Platform",
                   "Technical Solution Design for the integrated GRC solution")
    add_footer(doc, "GRC Platform")

    doc_control(doc)

    h1(doc, "Solution Overview", "2")
    para(doc,
         "The GRC Platform provides a unified system of record for governance, risk, compliance, "
         "asset management and third-party risk. It consolidates fragmented spreadsheets into a "
         "single multi-tenant application with strong access control, encryption, auditability "
         "and multilingual support.")
    h2(doc, "2.1  In-Scope Functional Modules")
    bullets(doc, [
        "Organization — Profile, Context, Business Processes, Business Impact Analysis (BIA).",
        "Compliance — Frameworks, Controls, Governance Documents, Evidence, Exceptions, KPIs.",
        "Risk Management — Risk Register, Assessment, Response, Risk-Control Matrix.",
        "Asset Management — Inventory, Classification.",
        "Third-Party Risk Management — Vendor inventory, Assessments, Monitoring, Issues/Follow-ups, Assessor workspace.",
    ])

    architecture(doc, 3, platform)
    tech_stack(doc, 4)

    # Module design
    h1(doc, "Module Design Detail", "5")
    h2(doc, "5.1  Organization")
    table(doc, ["Entity", "Key Attributes", "Notes"],
          [["Organization Profile", "Legal entity, structure, profile", "Tenant root context"],
           ["Context", "Internal/external context, interested parties", "ISO-aligned"],
           ["Process", "Name, owner, criticality", "Feeds BIA"],
           ["BIA", "RTO, RPO, dependencies, impact", "Linked to processes"]])
    h2(doc, "5.2  Compliance")
    table(doc, ["Entity", "Key Attributes", "Notes"],
          [["Framework", "Standard/regulation, version", "Multiple frameworks"],
           ["Control", "Reference, owner, status", "Mapped to frameworks"],
           ["Governance Document", "Policy/procedure, review cycle", "Approval workflow"],
           ["Evidence", "File (encrypted), linkage", "Bytes encrypted at rest"],
           ["Exception", "Reason, approver, expiry", "Waiver workflow"],
           ["KPI", "Metric, target, actual", "Dashboards"]])
    h2(doc, "5.3  Risk Management")
    table(doc, ["Entity", "Key Attributes", "Notes"],
          [["Risk", "Title, category, owner", "Central register"],
           ["Assessment", "Inherent/residual score", "Configurable methodology"],
           ["Response", "Treatment, owner, due date", "Tracked to closure"],
           ["Risk-Control Matrix", "Risk-control linkage", "Coverage analysis"]])
    h2(doc, "5.4  Asset Management")
    table(doc, ["Entity", "Key Attributes", "Notes"],
          [["Asset", "Name, type, owner", "Information/physical"],
           ["Classification", "Confidentiality, criticality", "Drives controls"]])
    h2(doc, "5.5  Third-Party Risk Management")
    table(doc, ["Entity", "Key Attributes", "Notes"],
          [["Vendor", "Profile, tier, owner", "Central inventory"],
           ["Assessment", "Questionnaire, score", "Due diligence"],
           ["Monitoring", "Status, signals", "Continuous"],
           ["Issue / Follow-up", "Finding, owner, due date", "Remediation"]])

    data_model_intro(doc, 6, platform)
    table(doc, ["Domain", "Representative Entities"],
          [["RBAC", "User, Role, Permission, CustomerAccount"],
           ["Organization", "OrganizationProfile, Context, Process, BIA"],
           ["Compliance", "Framework, Control, GovernanceDocument, Evidence, Exception, KPI"],
           ["Risk", "Risk, RiskAssessment, RiskResponse, RiskControlMatrix"],
           ["Assets", "Asset, AssetClassification"],
           ["TPRM", "Vendor, Assessment, Monitoring, Issue, FollowUp"],
           ["Translation", "DynamicTranslation (per record/field/locale)"]])

    security_section(doc, 7)
    integration_section(doc, 8, extra_rows=[
        ["GRC Reporting", "Export", "Risk register, compliance posture, KPI exports"],
    ])
    i18n_section(doc, 9)
    nfr_section(doc, 10)
    deployment_section(doc, 11)
    assumptions_risks(doc, 12)

    out = "E:/VSCode/GRC-AI/grc-app/docs/Solution-Design-GRC-Platform-Glimmora.docx"
    doc.save(out)
    return out


# =====================================================================
# DOCUMENT 2 — INTERNAL AUDIT PLATFORM
# =====================================================================
def build_audit():
    doc = Document(); style_base(doc)
    platform = "Internal Audit Platform"
    add_title_page(doc, "Internal Audit Platform",
                   "Technical Solution Design for the Internal Audit Management solution")
    add_footer(doc, "Internal Audit Platform")

    doc_control(doc)

    h1(doc, "Solution Overview", "2")
    para(doc,
         "The Internal Audit Platform digitises the full audit lifecycle — audit universe, "
         "risk-based planning, fieldwork, findings, CAPA and reporting — in a single multi-tenant "
         "application with role-based segregation across audit roles, encryption, and a complete "
         "audit trail.")
    h2(doc, "2.1  In-Scope Functional Modules")
    bullets(doc, [
        "Audit Universe & Risk Assessment — auditable entities, risk-based scoring, IA process setup.",
        "Audit Planning — annual/periodic plan, engagement scoping, resource allocation.",
        "Fieldwork & Execution — working papers, evidence, test procedures, auditee collaboration.",
        "Findings & CAPA — findings register, corrective/preventive actions, remediation workflow, escalation.",
        "Reporting — audit reports, dashboards, management & committee reporting.",
    ])

    architecture(doc, 3, platform)
    tech_stack(doc, 4)

    # Module design
    h1(doc, "Module Design Detail", "5")
    h2(doc, "5.1  Audit Universe & Risk Assessment")
    table(doc, ["Entity", "Key Attributes", "Notes"],
          [["Auditable Entity", "Name, process, unit", "Universe catalogue"],
           ["Risk Score", "Risk factors, rating", "Drives plan prioritisation"],
           ["IA Process", "Audit function structure", "Organization setup"]])
    h2(doc, "5.2  Audit Planning")
    table(doc, ["Entity", "Key Attributes", "Notes"],
          [["Audit Plan", "Period, entities, calendar", "Risk-based"],
           ["Engagement", "Objective, scope, criteria", "Per audit"],
           ["Resource Allocation", "Auditor assignment, schedule", "Capacity-aware"]])
    h2(doc, "5.3  Fieldwork & Execution")
    table(doc, ["Entity", "Key Attributes", "Notes"],
          [["Working Paper", "Procedure, result, conclusion", "Review/sign-off workflow"],
           ["Evidence", "File (encrypted), linkage", "Bytes encrypted at rest"],
           ["Test Procedure", "Steps, sample, outcome", "Standardised"],
           ["Auditee Request", "Information request, response", "Collaboration"]])
    h2(doc, "5.4  Findings & CAPA")
    table(doc, ["Entity", "Key Attributes", "Notes"],
          [["Finding", "Observation, rating, root cause", "Findings register"],
           ["CAPA", "Action, owner, due date, status", "Corrective/preventive"],
           ["Remediation", "Progress, verification", "Tracked to closure"]])
    h2(doc, "5.5  Reporting")
    table(doc, ["Entity", "Key Attributes", "Notes"],
          [["Audit Report", "Sections, findings, rating", "Engagement output"],
           ["Dashboard", "Plan status, findings, CAPA", "Real-time"]])

    data_model_intro(doc, 6, platform)
    table(doc, ["Domain", "Representative Entities"],
          [["RBAC", "User, Role (Audit Head/Manager/Auditor/Auditee), Permission, CustomerAccount"],
           ["Universe", "AuditableEntity, RiskScore, IAProcess"],
           ["Planning", "AuditPlan, Engagement, ResourceAllocation"],
           ["Fieldwork", "WorkingPaper, Evidence, TestProcedure"],
           ["Findings", "Finding, CAPA, Remediation"],
           ["Reporting", "AuditReport, Dashboard view"],
           ["Translation", "DynamicTranslation (per record/field/locale)"]])

    # Audit-specific RBAC detail
    h1(doc, "Audit Role Model (RBAC)", "7")
    table(doc, ["Role", "Access Summary"],
          [["Audit Head", "Full Internal Audit access; plan approval, all engagements."],
           ["Audit Manager", "Manage engagements, review working papers, findings."],
           ["Auditor", "Execute fieldwork, raise findings, draft reports."],
           ["Auditee", "Respond to requests, view/track assigned findings & CAPA."],
           ["Reviewer / Contributor", "Cross-module review/contribution as scoped."]])

    security_section(doc, 8)
    integration_section(doc, 9, extra_rows=[
        ["Audit Reporting", "Export", "Engagement reports, findings & CAPA exports"],
    ])
    i18n_section(doc, 10)
    nfr_section(doc, 11)
    deployment_section(doc, 12)
    assumptions_risks(doc, 13)

    out = "E:/VSCode/GRC-AI/grc-app/docs/Solution-Design-Internal-Audit-Platform-Glimmora.docx"
    doc.save(out)
    return out


if __name__ == "__main__":
    p1 = build_grc()
    p2 = build_audit()
    print("Created:")
    print(" -", p1)
    print(" -", p2)
