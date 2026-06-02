# -*- coding: utf-8 -*-
"""
Generate two Statement of Work (.docx) documents:
  1. GRC Platform
  2. Internal Audit Platform
Vendor: Glimmora International  |  Contact: info@glimmora.ai
"""
import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

VENDOR = "Glimmora International"
EMAIL = "info@glimmora.ai"
TODAY = datetime.date(2026, 6, 1).strftime("%d %B %Y")

BRAND = RGBColor(0x1F, 0x3A, 0x5F)      # deep navy
ACCENT = RGBColor(0x2E, 0x86, 0xC1)     # accent blue
GREY = RGBColor(0x55, 0x55, 0x55)


# ---------- helpers ----------
def set_cell_bg(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def style_base(doc):
    st = doc.styles["Normal"]
    st.font.name = "Calibri"
    st.font.size = Pt(10.5)
    st.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
    pf = st.paragraph_format
    pf.space_after = Pt(6)
    pf.line_spacing = 1.15


def add_title_page(doc, platform, subtitle):
    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(VENDOR)
    r.font.size = Pt(26)
    r.font.bold = True
    r.font.color.rgb = BRAND

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(EMAIL)
    r.font.size = Pt(11)
    r.font.color.rgb = ACCENT

    for _ in range(2):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("STATEMENT OF WORK")
    r.font.size = Pt(20)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0x22, 0x22, 0x22)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(platform)
    r.font.size = Pt(16)
    r.font.bold = True
    r.font.color.rgb = ACCENT

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(subtitle)
    r.font.size = Pt(11)
    r.font.italic = True
    r.font.color.rgb = GREY

    for _ in range(6):
        doc.add_paragraph()

    meta = [
        ("Document Title", f"Statement of Work — {platform}"),
        ("Prepared By", VENDOR),
        ("Contact", EMAIL),
        ("Document Date", TODAY),
        ("Version", "1.0"),
        ("Status", "Draft for Client Review"),
        ("Confidentiality", "Confidential — For Recipient Use Only"),
    ]
    t = doc.add_table(rows=0, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    for k, v in meta:
        row = t.add_row().cells
        row[0].width = Inches(2.2)
        row[1].width = Inches(3.8)
        rp = row[0].paragraphs[0].add_run(k)
        rp.bold = True
        rp.font.color.rgb = BRAND
        set_cell_bg(row[0], "EAF1F8")
        row[1].paragraphs[0].add_run(v)
    doc.add_page_break()


def h1(doc, text, n):
    p = doc.add_paragraph()
    p.space_before = Pt(10)
    r = p.add_run(f"{n}.  {text}")
    r.font.size = Pt(14)
    r.font.bold = True
    r.font.color.rgb = BRAND
    # bottom border
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), "2E86C1")
    pbdr.append(bottom)
    pPr.append(pbdr)
    return p


def h2(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(11.5)
    r.font.bold = True
    r.font.color.rgb = ACCENT
    return p


def para(doc, text):
    p = doc.add_paragraph(text)
    return p


def bullets(doc, items):
    for it in items:
        doc.add_paragraph(it, style="List Bullet")


def numbered(doc, items):
    for it in items:
        doc.add_paragraph(it, style="List Number")


def feature_table(doc, headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, htxt in enumerate(headers):
        set_cell_bg(hdr[i], "1F3A5F")
        run = hdr[i].paragraphs[0].add_run(htxt)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)
    for r in rows:
        cells = t.add_row().cells
        for i, val in enumerate(r):
            run = cells[i].paragraphs[0].add_run(val)
            run.font.size = Pt(9.5)
    return t


def signature_block(doc):
    h1(doc, "Acceptance & Authorisation", "13")
    para(doc,
         "By signing below, the parties acknowledge that they have read, understood and "
         "agreed to the terms, scope and deliverables described in this Statement of Work. "
         "This SOW becomes effective on the latest signature date below and is governed by "
         "the Master Services Agreement (or equivalent) executed between the parties.")
    t = doc.add_table(rows=1, cols=2)
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, label in enumerate([f"For and on behalf of {VENDOR} (Vendor)", "For and on behalf of the Client"]):
        set_cell_bg(hdr[i], "EAF1F8")
        run = hdr[i].paragraphs[0].add_run(label)
        run.bold = True
        run.font.color.rgb = BRAND
        run.font.size = Pt(10)
    fields = ["Name:", "Title:", "Signature:", "Date:"]
    for f in fields:
        cells = t.add_row().cells
        for i in range(2):
            cells[i].paragraphs[0].add_run("\n" + f).font.size = Pt(10)


def add_footer(doc, platform):
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"{VENDOR}  |  {EMAIL}  |  Statement of Work — {platform}  |  Confidential")
    r.font.size = Pt(8)
    r.font.color.rgb = GREY


# ---------- shared content sections ----------
def common_intro(doc, platform, intro_text):
    h1(doc, "Introduction & Background", "1")
    para(doc, intro_text)
    para(doc,
         f"This Statement of Work (\"SOW\") is entered into between {VENDOR} (\"the Vendor\", "
         f"\"we\", \"us\") and the Client (\"the Client\", \"you\"). It defines the scope of "
         f"work, deliverables, responsibilities, timeline and commercial framework for the "
         f"implementation of the {platform}. This SOW is to be read together with, and is "
         f"subject to, the Master Services Agreement or equivalent contractual arrangement "
         f"agreed between the parties.")

    h1(doc, "Objectives", "2")
    para(doc, "The engagement is designed to achieve the following outcomes:")


def commercials(doc, n):
    h1(doc, "Commercials & Payment Schedule", str(n))
    para(doc,
         "The commercial model below is indicative and to be finalised during contracting. "
         "All amounts are exclusive of applicable taxes, duties and third-party licensing "
         "costs unless explicitly stated.")
    feature_table(doc,
                  ["Milestone / Component", "Description", "Payment %"],
                  [
                      ["Project Initiation", "Mobilisation, kick-off, environment provisioning", "20%"],
                      ["Configuration & Build", "Module configuration, data model, RBAC setup", "30%"],
                      ["Data Migration & UAT", "Migration, integration, User Acceptance Testing", "25%"],
                      ["Go-Live", "Production cut-over and hypercare commencement", "15%"],
                      ["Closure & Sign-off", "Acceptance, handover, documentation", "10%"],
                  ])
    para(doc,
         "Recurring subscription / SaaS licensing, annual support & maintenance (AMC), and "
         "optional managed services are quoted separately in the commercial proposal.")


def assumptions(doc, n, extra=None):
    h1(doc, "Assumptions & Dependencies", str(n))
    base = [
        "The Client will nominate a Project Sponsor and a Single Point of Contact (SPOC) empowered to make decisions.",
        "The Client will provide timely access to subject-matter experts, source data, and required system credentials.",
        "Requirements signed off at the end of the design phase form the baseline; changes are handled via the Change Control process.",
        "Business and reference data will be provided in an agreed, structured electronic format.",
        "Third-party software, infrastructure or API licences required for integrations are procured by the Client unless stated otherwise.",
        "User Acceptance Testing will be completed by the Client within the agreed test window.",
        "Work is performed remotely unless on-site days are explicitly agreed and costed.",
    ]
    if extra:
        base.extend(extra)
    bullets(doc, base)


def out_of_scope(doc, n, items):
    h1(doc, "Out of Scope", str(n))
    para(doc, "Unless explicitly agreed in writing through Change Control, the following are excluded from this SOW:")
    bullets(doc, items)


def governance(doc, n):
    h1(doc, "Project Governance & Change Control", str(n))
    h2(doc, "Governance")
    bullets(doc, [
        "Weekly status reporting covering progress, risks, issues and decisions (RAID log).",
        "Steering Committee checkpoints at the end of each major phase.",
        "Escalation path defined between the Vendor and Client project managers.",
    ])
    h2(doc, "Change Control")
    para(doc,
         "Any change to scope, timeline or deliverables will be documented in a Change Request, "
         "assessed for impact on effort, cost and schedule, and implemented only after written "
         "approval by both parties.")


def roles_table(doc, n):
    h1(doc, "Roles & Responsibilities", str(n))
    feature_table(doc,
                  ["Role", "Party", "Responsibility"],
                  [
                      ["Project Sponsor", "Client", "Executive ownership, funding, escalation"],
                      ["Project Manager", "Both", "Day-to-day delivery, plan, RAID management"],
                      ["Solution / Implementation Lead", "Vendor", "Solution design, configuration, build"],
                      ["Business SMEs", "Client", "Requirements, validation, UAT sign-off"],
                      ["Data Owner", "Client", "Source data quality, migration approval"],
                      ["Technical / Infra Lead", "Both", "Environments, integrations, security"],
                      ["Training Lead", "Vendor", "End-user and administrator enablement"],
                  ])


def acceptance(doc, n):
    h1(doc, "Acceptance Criteria", str(n))
    para(doc, "A deliverable or milestone is deemed accepted when:")
    bullets(doc, [
        "It meets the agreed functional and non-functional requirements baseline.",
        "It successfully passes the agreed UAT test scripts with no open Severity-1 or Severity-2 defects.",
        "Supporting documentation (configuration, user guides, admin guides) has been handed over.",
        "Written sign-off is provided by the Client SPOC within the agreed review window (deemed accepted if no response within that window).",
    ])


def timeline(doc, n, phases):
    h1(doc, "Project Approach & Indicative Timeline", str(n))
    para(doc,
         "Delivery follows a phased, milestone-based methodology. The indicative durations below "
         "are refined during the Initiation phase and confirmed in the detailed project plan.")
    feature_table(doc, ["Phase", "Key Activities", "Indicative Duration"], phases)


# =====================================================================
# DOCUMENT 1 — GRC PLATFORM
# =====================================================================
def build_grc():
    doc = Document()
    style_base(doc)
    platform = "Governance, Risk & Compliance (GRC) Platform"
    add_title_page(doc, "GRC Platform",
                   "Implementation of an integrated Governance, Risk & Compliance solution")
    add_footer(doc, "GRC Platform")

    common_intro(doc, platform,
                 "Organisations today operate under increasing regulatory pressure and a complex "
                 "risk landscape. The GRC Platform delivers a single, integrated environment to "
                 "manage governance documentation, enterprise and operational risk, regulatory "
                 "compliance, asset management and third-party risk — replacing fragmented "
                 "spreadsheets and disconnected tools with a unified, auditable system of record.")
    bullets(doc, [
        "Establish a centralised, single source of truth for governance, risk and compliance data.",
        "Operationalise the organisation's risk management framework end-to-end.",
        "Map controls to one or more regulatory frameworks and track compliance posture in real time.",
        "Automate evidence collection, exceptions, KPIs and due-date reminders.",
        "Provide role-based, multi-tenant access with full audit traceability.",
        "Enable multilingual access (English, Arabic — RTL, and Latvian) across the platform.",
    ])

    # Scope / modules
    h1(doc, "Scope of Work — Functional Modules", "3")
    para(doc, "The following modules and capabilities are in scope for the GRC Platform implementation:")

    h2(doc, "3.1  Organization")
    feature_table(doc, ["Capability", "Description"],
                  [
                      ["Organization Profile", "Maintain legal entity, structure and business profile."],
                      ["Organizational Context", "Internal/external context, interested parties, scope statements."],
                      ["Business Processes", "Process inventory with ownership and criticality."],
                      ["Business Impact Analysis (BIA)", "Assess process criticality, RTO/RPO and dependencies."],
                  ])

    h2(doc, "3.2  Compliance")
    feature_table(doc, ["Capability", "Description"],
                  [
                      ["Frameworks", "Adopt and manage multiple regulatory/standard frameworks."],
                      ["Controls", "Control library mapped to frameworks with ownership and status."],
                      ["Governance Documents", "Policies/procedures with review cycles and approvals."],
                      ["Evidence Management", "Collect, store and link evidence to controls."],
                      ["Exceptions", "Raise, approve and track control exceptions/waivers."],
                      ["KPIs / Metrics", "Define and monitor compliance KPIs and dashboards."],
                  ])

    h2(doc, "3.3  Risk Management")
    feature_table(doc, ["Capability", "Description"],
                  [
                      ["Risk Register", "Central register of enterprise and operational risks."],
                      ["Risk Assessment", "Inherent/residual scoring with configurable methodology."],
                      ["Risk Response", "Treatment plans, owners, due dates and tracking."],
                      ["Risk-Control Matrix", "Link risks to mitigating controls for coverage analysis."],
                  ])

    h2(doc, "3.4  Asset Management")
    feature_table(doc, ["Capability", "Description"],
                  [
                      ["Asset Inventory", "Register of information and physical assets."],
                      ["Asset Classification", "Confidentiality/criticality classification and ownership."],
                  ])

    h2(doc, "3.5  Third-Party Risk Management (TPRM)")
    feature_table(doc, ["Capability", "Description"],
                  [
                      ["Vendor Inventory", "Central register of third parties and their profiles."],
                      ["Assessments", "Questionnaire-based vendor due-diligence assessments."],
                      ["Continuous Monitoring", "Ongoing monitoring of vendor risk posture."],
                      ["Issues & Follow-ups", "Track vendor issues, remediation and follow-ups."],
                      ["Assessor Workspace", "Dedicated assessor views for factory-style delivery."],
                  ])

    h2(doc, "3.6  Platform & Cross-Cutting Capabilities")
    bullets(doc, [
        "Role-Based Access Control (RBAC) with multi-tenant customer-account isolation.",
        "Configurable dashboards, reporting and exports.",
        "Automated notifications and daily due-date reminders.",
        "Field-level encryption at rest (AES-256-GCM) and TLS in transit for sensitive data.",
        "Full audit trail of create/edit/delete activity.",
        "Multilingual UI and dynamic data translation (English, Arabic RTL, Latvian).",
        "AI-assisted guidance via integrated assistant (where enabled).",
    ])

    # Deliverables
    h1(doc, "Deliverables", "4")
    feature_table(doc, ["#", "Deliverable", "Format"],
                  [
                      ["D1", "Solution Design & Requirements Baseline document", "Document"],
                      ["D2", "Configured GRC environment (all in-scope modules)", "Working system"],
                      ["D3", "RBAC role matrix and tenant configuration", "Configuration + doc"],
                      ["D4", "Data migration of agreed master/reference data", "Loaded data + report"],
                      ["D5", "Integration setup (as agreed in design)", "Working integration"],
                      ["D6", "UAT support pack and defect log", "Documents"],
                      ["D7", "Administrator & end-user guides", "Documents"],
                      ["D8", "Training sessions (admin + end-user)", "Sessions"],
                      ["D9", "Go-live and hypercare support", "Service"],
                  ])

    timeline(doc, 5, [
        ["1. Initiation", "Kick-off, governance setup, environment provisioning", "1–2 weeks"],
        ["2. Design", "Requirements workshops, solution design, sign-off", "2–4 weeks"],
        ["3. Build & Configure", "Module configuration, RBAC, dashboards, integrations", "4–8 weeks"],
        ["4. Data & Test", "Data migration, SIT, UAT support", "2–4 weeks"],
        ["5. Deploy", "Production cut-over, go-live", "1 week"],
        ["6. Hypercare", "Post-go-live stabilisation support", "2–4 weeks"],
    ])

    roles_table(doc, 6)
    acceptance(doc, 7)
    assumptions(doc, 8, extra=[
        "The number of regulatory frameworks and controls to be loaded is agreed during design.",
        "TPRM assessment questionnaires are provided or selected from standard templates.",
    ])
    out_of_scope(doc, 9, [
        "Bespoke module development beyond configuration of the standard platform.",
        "Custom integrations not identified and agreed during the Design phase.",
        "Content authoring of the Client's policies, risks or controls (the Client provides content).",
        "Independent regulatory certification or audit attestation services.",
        "Ongoing managed services beyond the hypercare period (available separately).",
    ])
    governance(doc, 10)
    commercials(doc, 11)

    h1(doc, "Support & Maintenance", "12")
    para(doc,
         "Following hypercare, ongoing support is available under a separate Annual Maintenance "
         "Contract (AMC) covering platform updates, defect resolution, and service-desk support "
         "under agreed SLAs. Tiered support options can be tailored to the Client's needs.")

    signature_block(doc)

    out = "E:/VSCode/GRC-AI/grc-app/docs/SOW-GRC-Platform-Glimmora.docx"
    doc.save(out)
    return out


# =====================================================================
# DOCUMENT 2 — INTERNAL AUDIT PLATFORM
# =====================================================================
def build_audit():
    doc = Document()
    style_base(doc)
    platform = "Internal Audit Platform"
    add_title_page(doc, "Internal Audit Platform",
                   "Implementation of an end-to-end Internal Audit Management solution")
    add_footer(doc, "Internal Audit Platform")

    common_intro(doc, platform,
                 "The Internal Audit Platform digitises the complete internal audit lifecycle — "
                 "from risk-based audit planning through fieldwork, findings, corrective actions "
                 "and reporting. It replaces manual, spreadsheet-driven audit processes with a "
                 "structured, traceable and collaborative workflow that strengthens assurance, "
                 "improves auditor productivity, and gives management real-time visibility of the "
                 "audit universe and remediation status.")
    bullets(doc, [
        "Build and maintain a risk-based audit universe and annual audit plan.",
        "Standardise audit execution — planning, fieldwork, working papers and evidence.",
        "Capture findings and drive Corrective and Preventive Action (CAPA) to closure.",
        "Provide real-time dashboards on audit progress, findings and overdue actions.",
        "Enforce role-based segregation across audit roles (Head, Manager, Auditor, Auditee).",
        "Maintain a defensible, time-stamped audit trail for quality assurance and external review.",
    ])

    # Scope / modules
    h1(doc, "Scope of Work — Functional Modules", "3")
    para(doc, "The following modules and capabilities are in scope for the Internal Audit Platform implementation:")

    h2(doc, "3.1  Audit Universe & Risk Assessment")
    feature_table(doc, ["Capability", "Description"],
                  [
                      ["Audit Universe", "Catalogue of auditable entities, processes and units."],
                      ["Risk-Based Scoring", "Score and rank auditable entities by risk."],
                      ["Audit Organization", "Audit function structure, including IA process setup."],
                  ])

    h2(doc, "3.2  Audit Planning")
    feature_table(doc, ["Capability", "Description"],
                  [
                      ["Annual / Periodic Plan", "Build the risk-based audit plan and calendar."],
                      ["Engagement Scoping", "Define objectives, scope, criteria and resourcing."],
                      ["Resource Allocation", "Assign auditors and schedule engagements."],
                  ])

    h2(doc, "3.3  Fieldwork & Execution")
    feature_table(doc, ["Capability", "Description"],
                  [
                      ["Working Papers", "Structured workpapers with review/sign-off workflow."],
                      ["Evidence Management", "Attach, link and secure audit evidence."],
                      ["Test Procedures", "Document test steps, results and conclusions."],
                      ["Auditee Collaboration", "Request information and responses from auditees."],
                  ])

    h2(doc, "3.4  Findings & CAPA")
    feature_table(doc, ["Capability", "Description"],
                  [
                      ["Findings Register", "Record observations, ratings, root cause and impact."],
                      ["CAPA Tracking", "Corrective/Preventive actions with owners and due dates."],
                      ["Remediation Workflow", "Track action status through to verified closure."],
                      ["Overdue Escalation", "Automated reminders and escalation for overdue actions."],
                  ])

    h2(doc, "3.5  Reporting")
    feature_table(doc, ["Capability", "Description"],
                  [
                      ["Audit Reports", "Generate structured engagement reports."],
                      ["Dashboards", "Real-time status of plan, findings and CAPA."],
                      ["Management & Committee Reporting", "Summarised assurance views for leadership/Audit Committee."],
                  ])

    h2(doc, "3.6  Platform & Cross-Cutting Capabilities")
    bullets(doc, [
        "Role-Based Access Control with dedicated audit roles (Audit Head, Audit Manager, Auditor, Auditee) and multi-tenant isolation.",
        "Configurable dashboards, exports and report generation.",
        "Automated notifications and daily due-date reminders for findings and CAPA.",
        "Field-level encryption at rest (AES-256-GCM) and TLS in transit for sensitive audit data.",
        "Complete, time-stamped audit trail across the engagement lifecycle.",
        "Multilingual UI and dynamic data translation (English, Arabic RTL, Latvian).",
        "AI-assisted guidance via integrated assistant (where enabled).",
    ])

    # Deliverables
    h1(doc, "Deliverables", "4")
    feature_table(doc, ["#", "Deliverable", "Format"],
                  [
                      ["D1", "Solution Design & Requirements Baseline document", "Document"],
                      ["D2", "Configured Internal Audit environment (all in-scope modules)", "Working system"],
                      ["D3", "Audit role matrix and tenant configuration", "Configuration + doc"],
                      ["D4", "Audit universe and plan setup with agreed data", "Loaded data + report"],
                      ["D5", "Working-paper and report templates", "Templates"],
                      ["D6", "UAT support pack and defect log", "Documents"],
                      ["D7", "Administrator & auditor user guides", "Documents"],
                      ["D8", "Training sessions (administrators, auditors, auditees)", "Sessions"],
                      ["D9", "Go-live and hypercare support", "Service"],
                  ])

    timeline(doc, 5, [
        ["1. Initiation", "Kick-off, governance setup, environment provisioning", "1–2 weeks"],
        ["2. Design", "Audit methodology workshops, solution design, sign-off", "2–4 weeks"],
        ["3. Build & Configure", "Universe, plan, workpaper & report templates, RBAC", "3–6 weeks"],
        ["4. Data & Test", "Data setup, SIT, UAT support", "2–3 weeks"],
        ["5. Deploy", "Production cut-over, go-live", "1 week"],
        ["6. Hypercare", "Post-go-live stabilisation support", "2–4 weeks"],
    ])

    roles_table(doc, 6)
    acceptance(doc, 7)
    assumptions(doc, 8, extra=[
        "The Client's audit methodology and rating scales are provided during the Design phase.",
        "Working-paper and report templates are agreed and frozen before configuration.",
    ])
    out_of_scope(doc, 9, [
        "Performing actual internal audit engagements on behalf of the Client.",
        "Bespoke module development beyond configuration of the standard platform.",
        "Custom integrations not identified and agreed during the Design phase.",
        "Authoring of the Client's audit programmes or finding content.",
        "External audit, certification or independent assurance services.",
        "Ongoing managed services beyond the hypercare period (available separately).",
    ])
    governance(doc, 10)
    commercials(doc, 11)

    h1(doc, "Support & Maintenance", "12")
    para(doc,
         "Following hypercare, ongoing support is available under a separate Annual Maintenance "
         "Contract (AMC) covering platform updates, defect resolution, and service-desk support "
         "under agreed SLAs. Tiered support options can be tailored to the Client's needs.")

    signature_block(doc)

    out = "E:/VSCode/GRC-AI/grc-app/docs/SOW-Internal-Audit-Platform-Glimmora.docx"
    doc.save(out)
    return out


if __name__ == "__main__":
    p1 = build_grc()
    p2 = build_audit()
    print("Created:")
    print(" -", p1)
    print(" -", p2)
