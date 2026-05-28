"""
Generate platform documentation .docx files for Internal Audit and GRC.

Produces 8 documents in docs/platform-docs/:
  Internal Audit:  Target Audience, Technical Documentation, PPT, USP
  GRC:             Target Audience, Technical Documentation, PPT, USP

Run:  python scripts/generate-platform-docs.py
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ---------- Styling helpers ----------

BRAND_PRIMARY = RGBColor(0x0B, 0x3D, 0x91)   # deep blue
BRAND_ACCENT  = RGBColor(0x1E, 0x88, 0xE5)   # sky blue
TEXT_DARK     = RGBColor(0x1F, 0x2A, 0x44)
TEXT_MUTED    = RGBColor(0x55, 0x5F, 0x77)
TABLE_HEADER_BG = "0B3D91"
TABLE_ALT_BG    = "F2F5FB"


def set_cell_bg(cell, hex_color: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def add_horizontal_rule(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "1E88E5")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def configure_styles(doc: Document):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = TEXT_DARK

    for h_name, size, color in [
        ("Heading 1", 22, BRAND_PRIMARY),
        ("Heading 2", 16, BRAND_PRIMARY),
        ("Heading 3", 13, BRAND_ACCENT),
    ]:
        s = styles[h_name]
        s.font.name = "Calibri"
        s.font.size = Pt(size)
        s.font.bold = True
        s.font.color.rgb = color


def add_cover(doc: Document, platform: str, doc_type: str, subtitle: str):
    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("\n\n\n")  # vertical spacing

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(platform.upper())
    r.font.size = Pt(14)
    r.font.bold = True
    r.font.color.rgb = BRAND_ACCENT

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(doc_type)
    r.font.size = Pt(36)
    r.font.bold = True
    r.font.color.rgb = BRAND_PRIMARY

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_horizontal_rule(p)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(subtitle)
    r.font.size = Pt(14)
    r.font.italic = True
    r.font.color.rgb = TEXT_MUTED

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("\n\nGlimmora International\ninfo@glimmora.ai")
    r.font.size = Pt(11)
    r.font.color.rgb = TEXT_MUTED

    doc.add_page_break()


def add_section_heading(doc: Document, text: str, level: int = 1):
    p = doc.add_heading(text, level=level)
    return p


def add_paragraph(doc: Document, text: str, bold: bool = False, italic: bool = False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.bold = bold
    r.font.italic = italic
    return p


def add_bullets(doc: Document, items: list):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_numbered(doc: Document, items: list):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)


def add_table(doc: Document, headers: list, rows: list, col_widths: list = None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        r = p.add_run(h)
        r.font.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r.font.size = Pt(11)
        set_cell_bg(hdr[i], TABLE_HEADER_BG)
        hdr[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for idx, row_data in enumerate(rows):
        row = table.add_row().cells
        for i, v in enumerate(row_data):
            row[i].text = ""
            p = row[i].paragraphs[0]
            r = p.add_run(str(v))
            r.font.size = Pt(10.5)
            if idx % 2 == 1:
                set_cell_bg(row[i], TABLE_ALT_BG)

    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)

    return table


def add_callout(doc: Document, title: str, body: str):
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    set_cell_bg(cell, TABLE_ALT_BG)
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(title + "\n")
    r.font.bold = True
    r.font.color.rgb = BRAND_PRIMARY
    r.font.size = Pt(11)
    r2 = p.add_run(body)
    r2.font.size = Pt(10.5)
    r2.font.color.rgb = TEXT_DARK


# ============================================================
# CONTENT — INTERNAL AUDIT
# ============================================================

def build_ia_target_audience(doc: Document):
    add_cover(doc, "Internal Audit Platform", "Target Audience",
              "Who this platform serves and the value it delivers to each persona")

    add_section_heading(doc, "1. Purpose of This Document", 1)
    add_paragraph(doc,
        "This document identifies the primary and secondary audiences for the Glimmora "
        "Internal Audit Platform, the problems each audience faces today, and the "
        "specific capabilities the platform delivers to address those problems. It is "
        "intended for solution buyers, internal champions, sales engineers, and "
        "onboarding teams who need to map platform features to stakeholder value.")

    add_section_heading(doc, "2. Primary Audiences", 1)

    add_section_heading(doc, "2.1 Chief Audit Executive (CAE) / Audit Head", 2)
    add_paragraph(doc,
        "Owns the annual audit plan, reports to the Audit Committee, and is "
        "accountable for the independence and effectiveness of the internal audit "
        "function.")
    add_paragraph(doc, "Key pain points addressed:", bold=True)
    add_bullets(doc, [
        "Fragmented audit data across spreadsheets and email threads",
        "Inability to demonstrate risk-based coverage to the Audit Committee",
        "Manual consolidation of CAPA status before every board meeting",
        "Lack of real-time visibility into engagement status",
    ])
    add_paragraph(doc, "Platform value:", bold=True)
    add_bullets(doc, [
        "Dashboard showing audit-plan execution, open findings by rating, and CAPA aging",
        "Risk-based audit universe with AI-recommended engagements",
        "One-click board-ready audit reports with charts and findings register",
        "Configurable scoring ranges, risk factors, and impact/probability matrices",
    ])

    add_section_heading(doc, "2.2 Audit Manager", 2)
    add_paragraph(doc,
        "Plans engagements, allocates auditors, reviews workpapers, and signs off on "
        "findings before they reach the Audit Head.")
    add_paragraph(doc, "Key pain points addressed:", bold=True)
    add_bullets(doc, [
        "Workpaper review consumes evenings and weekends",
        "Evidence requests get lost in email; auditees miss deadlines",
        "No standardised template for findings — quality varies by auditor",
        "Difficult to track team workload and reassign engagements mid-cycle",
    ])
    add_paragraph(doc, "Platform value:", bold=True)
    add_bullets(doc, [
        "Workpaper module with version control, AI-assisted review, and inline comments",
        "Evidence Request workflow with auto-reminders and escalation to auditees",
        "Standardised finding templates with mandatory fields (root cause, recommendation, severity)",
        "Auditor schedule view showing capacity, in-flight engagements, and overdue tasks",
    ])

    add_section_heading(doc, "2.3 Auditor (Field Auditor / Senior Auditor)", 2)
    add_paragraph(doc,
        "Executes engagements end-to-end: planning meetings, fieldwork, evidence "
        "collection, control testing, and drafting findings.")
    add_paragraph(doc, "Key pain points addressed:", bold=True)
    add_bullets(doc, [
        "Re-keying the same process information across multiple audits",
        "Chasing evidence from auditees by email and phone",
        "Manual creation of audit programs and test scripts for every engagement",
        "Hard to surface historical findings on the same process",
    ])
    add_paragraph(doc, "Platform value:", bold=True)
    add_bullets(doc, [
        "Reusable Process Universe — define a process once, reuse across engagements",
        "Self-service evidence portal for auditees with deadline tracking",
        "AI-generated workpapers based on process risks and historical audits",
        "Document Library with full-text search across past engagements",
    ])

    add_section_heading(doc, "2.4 Auditee (Process / Control Owner)", 2)
    add_paragraph(doc,
        "Receives evidence requests, responds to findings, and owns Corrective and "
        "Preventive Actions (CAPA) for issues identified in their area.")
    add_paragraph(doc, "Key pain points addressed:", bold=True)
    add_bullets(doc, [
        "Evidence requests arrive without context — what is being tested and why?",
        "No single place to see all open audit actions assigned to them",
        "Repeated requests for the same documents across audits",
        "No visibility into how their CAPA progress is being tracked",
    ])
    add_paragraph(doc, "Platform value:", bold=True)
    add_bullets(doc, [
        "Auditee Dashboard listing open requests, findings, and CAPA actions in one place",
        "Each request shows the related control, test objective, and acceptance criteria",
        "Reusable evidence library — submit once, link to multiple engagements",
        "CAPA progress tracker with attachment uploads and AI-review for closure quality",
    ])

    add_section_heading(doc, "3. Secondary Audiences", 1)

    add_section_heading(doc, "3.1 Audit Committee / Board", 2)
    add_paragraph(doc,
        "Consumes audit outcomes; does not operate the platform daily but relies on "
        "its outputs to discharge governance responsibilities.")
    add_bullets(doc, [
        "Board-ready reports generated directly from platform data",
        "Heat-maps of risk coverage and CAPA closure trends",
        "Independent visibility into findings without going through management",
    ])

    add_section_heading(doc, "3.2 External Auditors / Regulators", 2)
    add_paragraph(doc,
        "Periodically request access to workpapers, evidence, and CAPA records during "
        "statutory audits or regulatory inspections.")
    add_bullets(doc, [
        "Read-only roles can be granted with time-boxed access",
        "Full audit trail of who viewed, edited, or approved each artefact",
        "Export packs (PDF/DOCX) for offline review",
    ])

    add_section_heading(doc, "3.3 IT / Information Security Team", 2)
    add_paragraph(doc,
        "Responsible for the secure operation of the platform and integration with "
        "enterprise identity and storage systems.")
    add_bullets(doc, [
        "AES-256-GCM encryption at rest, TLS 1.2+ in transit",
        "Role-based access control with department-scoped roles",
        "Detailed activity logs feeding into the SIEM",
    ])

    add_section_heading(doc, "4. Target Industries", 1)
    add_table(doc,
        ["Industry", "Why Internal Audit Matters", "Specific Drivers"],
        [
            ["Banking & Financial Services", "Highly regulated; mandatory internal audit function", "RBI guidelines, SOX-equivalent, IFC compliance"],
            ["Insurance", "Solvency, claims integrity, IRDAI oversight", "Risk-based capital, ORSA"],
            ["Healthcare", "Patient safety, billing integrity, NABH/JCI accreditation", "HIPAA-like local laws, accreditation cycles"],
            ["Manufacturing", "Process safety, inventory accuracy, supplier audits", "ISO 9001, supplier code of conduct"],
            ["IT / ITeS", "Information security, data privacy, client SLAs", "ISO 27001, SOC 2, client audit clauses"],
            ["Public Sector / PSUs", "Compliance with CAG, internal vigilance", "Mandatory IA charter, CVC guidelines"],
            ["Retail & E-commerce", "Cash management, shrinkage, vendor compliance", "Loss prevention, vendor due-diligence"],
        ],
        col_widths=[4.5, 5.5, 6.0])

    add_section_heading(doc, "5. Buying Personas", 1)
    add_table(doc,
        ["Persona", "Role in Decision", "What They Want to Hear"],
        [
            ["Chief Audit Executive", "Economic buyer", "Reduced audit cycle time, better Board reporting, defensible methodology"],
            ["CFO / Finance Director", "Budget approver", "TCO, ROI, reduction in external audit fees"],
            ["Chief Risk Officer", "Influencer", "Integration with risk register, shared taxonomy"],
            ["Head of IT / CISO", "Technical gatekeeper", "Security posture, deployment options, data residency"],
            ["Audit Manager", "Day-to-day champion", "Productivity gains, less manual rework"],
        ],
        col_widths=[4.5, 4.5, 7.0])

    add_section_heading(doc, "6. Out of Scope (Not Our Audience)", 1)
    add_paragraph(doc,
        "The Internal Audit Platform is purpose-built for second- and third-line "
        "assurance teams. The following audiences are explicitly not the focus:")
    add_bullets(doc, [
        "External / statutory auditors performing financial-statement audits as a primary workflow (we serve them as consumers of internal audit output)",
        "Single-user freelance auditors looking for a personal productivity tool",
        "Organisations without a formal internal audit function or charter",
    ])


def build_ia_technical(doc: Document):
    add_cover(doc, "Internal Audit Platform", "Technical Documentation",
              "Architecture, technology stack, security model, and integration points")

    add_section_heading(doc, "1. Solution Overview", 1)
    add_paragraph(doc,
        "The Internal Audit Platform is a web-based, multi-tenant SaaS application "
        "built on the modern Next.js 16 App Router stack. It provides end-to-end "
        "internal audit lifecycle management — from risk-based annual planning "
        "through engagement execution, finding management, and CAPA closure — and is "
        "augmented by a dedicated Python AI backend for document understanding, "
        "workpaper generation, and natural-language query.")

    add_section_heading(doc, "2. Architecture", 1)
    add_section_heading(doc, "2.1 High-Level Components", 2)
    add_table(doc,
        ["Layer", "Technology", "Responsibility"],
        [
            ["Presentation", "Next.js 16 (App Router), React 19, Tailwind CSS, shadcn/ui", "Server components, route groups, RBAC-filtered navigation"],
            ["API", "Next.js Route Handlers + NextAuth v5 (JWT)", "REST endpoints, session validation, permission enforcement"],
            ["Business Logic", "TypeScript service modules under src/lib", "Audit workflows, scoring, escalation, notifications"],
            ["Persistence", "Prisma 6 ORM + PostgreSQL (Neon / DO Managed)", "Schema-first models, migrations, encrypted fields"],
            ["AI Services", "Python FastAPI backend + MongoDB", "Document ingestion, embeddings, AI workpapers, audit Q&A"],
            ["File Storage", "Object storage (DO Spaces / S3-compatible)", "Workpapers, evidence, attachments — encrypted at rest"],
            ["Email", "SMTP relay (configurable)", "Notifications, evidence requests, escalations"],
        ],
        col_widths=[3.5, 5.0, 7.5])

    add_section_heading(doc, "2.2 Deployment Topology", 2)
    add_paragraph(doc,
        "The reference deployment runs on DigitalOcean App Platform. A single git "
        "branch (GRC-MultiTenant) is wired to two App Platform apps — UAT and "
        "Production — each with its own database and environment variables. Risky "
        "changes are gated by per-app feature flags so they can be validated in UAT "
        "for 24–48 hours before being enabled in Production.")

    add_section_heading(doc, "3. Module Inventory", 1)
    add_table(doc,
        ["Module", "Capabilities"],
        [
            ["Dashboard", "Audit-plan progress, CAPA aging, risk-by-rating, auditor schedule, drill-downs"],
            ["Audit Universe", "Auditable entities, process/department mapping, last-audited tracking"],
            ["Risk Universe & Register", "Risk taxonomy, residual scoring, AI-recommended audits, import/export"],
            ["Audit Planning", "Annual plan, engagement creation, resource allocation, AI plan generation"],
            ["Fieldwork", "Engagement workspace, workpapers (AI + manual), evidence requests, findings"],
            ["Findings & CAPA", "Finding lifecycle, severity, root cause, CAPA assignment, AI closure review"],
            ["Reports", "Engagement reports, board reports, document downloads (PDF/DOCX)"],
            ["Document Library", "Full-text search across past artefacts, AI ingestion, recent searches"],
            ["Settings", "Audit types, categories, processes, scoring config, escalation, user management"],
        ],
        col_widths=[4.5, 11.0])

    add_section_heading(doc, "4. Authentication & Authorization", 1)
    add_paragraph(doc, "NextAuth v5 with JWT sessions. Role-Based Access Control is enforced at three layers:")
    add_bullets(doc, [
        "Navigation: src/lib/navigation.ts filters menu items by resource:action permission",
        "API: withAuth wrapper in src/lib/api-auth.ts validates session + permission per route",
        "UI: usePermissions / useHasRole hooks gate buttons, fields, and entire pages",
    ])
    add_paragraph(doc, "Audit-relevant roles:", bold=True)
    add_table(doc,
        ["Role", "Scope", "Typical Use"],
        [
            ["AuditHead", "Full Internal Audit access across organisation", "Owns the plan, signs off engagements"],
            ["AuditManager", "Create/edit engagements, review workpapers", "Day-to-day execution oversight"],
            ["Auditor", "Execute assigned engagements", "Fieldwork, evidence collection, draft findings"],
            ["Auditee", "Respond to requests assigned to them", "Process / control owners"],
            ["Reviewer", "Cross-module read + comment", "Stakeholder visibility without edit rights"],
            ["DepartmentReviewer", "Department-scoped read", "Functional heads"],
        ],
        col_widths=[4.5, 5.5, 6.0])

    add_section_heading(doc, "5. Data Model (Selected Entities)", 1)
    add_bullets(doc, [
        "AuditUniverse — entities and processes eligible for audit",
        "Risk, RiskFactor, ImpactScale, ProbabilityScale, ScoringRange",
        "Engagement — annual-plan instance with status, scope, team",
        "Fieldwork — execution workspace child of Engagement",
        "Workpaper, AiWorkpaper, EvidenceRequest, Attachment",
        "Finding — outcome of testing, links to Control and Process",
        "Capa — corrective/preventive actions tied to Findings",
        "AuditReport — generated engagement and management reports",
    ])

    add_section_heading(doc, "6. AI Services", 1)
    add_paragraph(doc,
        "The Python AI backend is deployed as an independent service with its own "
        "MongoDB datastore. The Next.js app communicates with it over HTTPS using a "
        "shared secret (PYTHON_API_SECRET).")
    add_table(doc,
        ["Capability", "Endpoint / Module", "Description"],
        [
            ["Document Ingestion", "/documents/ingest", "OCR, chunking, embeddings stored in MongoDB"],
            ["Audit Q&A", "/fieldwork/audit-query", "RAG over ingested documents"],
            ["AI Workpapers", "/fieldwork/ai-workpapers/generate", "Generate workpaper drafts per process / risk"],
            ["CAPA AI Review", "/capa/ai-review", "Assesses closure-evidence sufficiency"],
            ["Risk Identification", "/risks/suggest", "Suggests risks from process descriptions"],
            ["AI Audit Plan", "/audit-planning/from-ai", "Recommends engagements from risk register"],
            ["Translations", "/translate", "Dynamic translation of user-entered data"],
        ],
        col_widths=[4.0, 5.5, 6.5])

    add_section_heading(doc, "7. Security", 1)
    add_bullets(doc, [
        "TLS 1.2+ enforced for all inbound and outbound traffic",
        "AES-256-GCM field-level encryption for sensitive Bytes columns (see docs/SECURITY.md)",
        "Encryption kill switch via ENCRYPTION_ENABLED env var per app",
        "Secrets stored as encrypted env vars on DigitalOcean — never in git",
        "Password hashing with bcrypt; configurable password policies",
        "Activity log captures who/what/when for every state-changing API",
        "Safe-log helper (src/lib/safe-log.ts) redacts known-sensitive keys before logging",
    ])
    add_callout(doc, "Compliance Mapping",
        "Controls map directly to ISO 27001 A.10 (Cryptography), A.9 (Access Control), "
        "A.12 (Operations Security); SOC 2 CC6 (Logical Access) and CC7 (System "
        "Operations); GDPR Art. 32 (Security of Processing); and PCI DSS Req. 3 & 4.")

    add_section_heading(doc, "8. Integration Points", 1)
    add_table(doc,
        ["Integration", "Direction", "Notes"],
        [
            ["SSO (SAML / OIDC)", "Inbound", "Optional; NextAuth supports custom providers"],
            ["SMTP", "Outbound", "Notifications, reminders, escalations"],
            ["Object Storage (S3-compatible)", "Bi-directional", "Workpaper and evidence storage"],
            ["Webhook (export to ITSM)", "Outbound", "Findings/CAPA push to ServiceNow / Jira (roadmap)"],
            ["SIEM", "Outbound", "Activity log streaming"],
        ],
        col_widths=[4.5, 3.5, 8.0])

    add_section_heading(doc, "9. Scheduled Jobs", 1)
    add_bullets(doc, [
        "Daily 08:00 UTC — /api/cron/due-reminders sends notifications for items due within 24 hours (evidence, CAPA, policy reviews)",
        "Configurable escalation rules per finding severity",
    ])

    add_section_heading(doc, "10. Internationalisation", 1)
    add_paragraph(doc,
        "The platform is shipped with English, Arabic (RTL), and Latvian. Static UI "
        "strings are translated through a phrase-based t() function; user-entered data "
        "is translated dynamically by the Python backend and cached in the "
        "DynamicTranslation table on create/edit only.")

    add_section_heading(doc, "11. Non-Functional Targets", 1)
    add_table(doc,
        ["Aspect", "Target"],
        [
            ["Page load (P95)", "< 2.0 seconds on broadband"],
            ["API response (P95)", "< 600 ms for read, < 1.5 s for write"],
            ["Concurrent users per tenant", "500+ on standard tier"],
            ["Uptime SLA", "99.5% on UAT, 99.9% on Production"],
            ["RPO / RTO", "RPO 24 h (daily backup), RTO 4 h"],
            ["Browser support", "Chrome, Edge, Firefox, Safari — current and prior major versions"],
        ],
        col_widths=[5.0, 10.5])

    add_section_heading(doc, "12. Glossary", 1)
    add_table(doc,
        ["Term", "Meaning"],
        [
            ["CAPA", "Corrective and Preventive Action"],
            ["Engagement", "A discrete internal audit assignment"],
            ["Workpaper", "Evidence-bearing document created during fieldwork"],
            ["Risk Universe", "Master catalogue of risks the organisation tracks"],
            ["Audit Universe", "Master catalogue of auditable entities and processes"],
            ["RBAC", "Role-Based Access Control"],
        ],
        col_widths=[3.5, 12.0])


def build_ia_ppt(doc: Document):
    # PPT-style: each H1 = a slide title; bullets = slide body
    add_cover(doc, "Internal Audit Platform", "Presentation Deck",
              "Slide-by-slide narrative for stakeholder briefings")

    slides = [
        ("Slide 1 — The Problem",
         ["Internal audit is buried in spreadsheets, email threads, and tribal knowledge",
          "Audit Committee asks: 'Are we covering the right risks?' — answer takes weeks",
          "CAPA closure is opaque; findings repeat year after year",
          "External audit fees rise because the internal trail is hard to follow"]),
        ("Slide 2 — Why Now",
         ["Regulators demand demonstrable, risk-based assurance",
          "Boards expect dashboard-grade reporting, not PDFs",
          "AI has matured to the point where workpapers, evidence review, and Q&A can be automated safely",
          "Hybrid workforces need a single source of truth that is not a shared drive"]),
        ("Slide 3 — Our Solution",
         ["End-to-end Internal Audit lifecycle on one platform",
          "Risk-based planning  ->  fieldwork  ->  findings  ->  CAPA  ->  reporting",
          "AI-augmented at every stage — but the human remains in control",
          "Built for multi-tenant SaaS; deployable as single-tenant if required"]),
        ("Slide 4 — Audit Lifecycle in One View",
         ["Audit Universe and Risk Universe feed the Annual Plan",
          "Engagements drive Fieldwork; Fieldwork generates Findings",
          "Findings drive CAPA; CAPA closures feed Reports",
          "Reports go straight to the Audit Committee — no manual assembly"]),
        ("Slide 5 — Module Map",
         ["Dashboard — execution + CAPA aging + risk-by-rating",
          "Audit Universe + Risk Universe — your taxonomies, your scoring",
          "Audit Planning — manual or AI-generated annual plan",
          "Fieldwork — workpapers, evidence requests, findings",
          "CAPA Tracking — owner, due date, AI-review of closure quality",
          "Reports + Document Library — board-ready outputs and searchable history"]),
        ("Slide 6 — AI Capabilities",
         ["AI-Recommended Audits — surface where to audit next",
          "AI Workpapers — draft tests against process and risk",
          "Audit Q&A — chat with past engagements and evidence",
          "AI CAPA Review — flags weak or off-topic closure evidence",
          "Risk Suggestions — propose risks from process narratives",
          "Translation — Arabic and Latvian for global teams"]),
        ("Slide 7 — Security & Trust",
         ["AES-256-GCM at rest, TLS 1.2+ in transit",
          "RBAC with 8+ audit-specific roles",
          "Full activity log — every read, edit, approval",
          "Maps to ISO 27001, SOC 2, GDPR, PCI",
          "Kill switch + 90-day master-key rotation"]),
        ("Slide 8 — Deployment Options",
         ["Multi-tenant SaaS on DigitalOcean (default)",
          "Single-tenant on customer cloud (AWS / Azure / GCP)",
          "Hybrid — app in customer cloud, AI services on Glimmora backbone",
          "UAT + Production environments included, gated by feature flags"]),
        ("Slide 9 — Stakeholder Value",
         ["Audit Head — defensible methodology and board-ready reporting",
          "Audit Manager — workpaper review time cut in half",
          "Auditor — reusable processes, AI workpaper drafts, self-service evidence",
          "Auditee — one dashboard for every request, finding, and CAPA",
          "CFO — reduced external audit fees and TCO"]),
        ("Slide 10 — Implementation Approach",
         ["Week 1–2: Discovery — map audit universe and risk taxonomy",
          "Week 3–4: Configuration — scoring, roles, processes, departments",
          "Week 5–6: Pilot engagement with one audit team",
          "Week 7–8: Migration of historical findings + CAPA",
          "Week 9+: Full rollout with weekly office hours for 90 days"]),
        ("Slide 11 — Pricing Model",
         ["Per-tenant annual subscription tied to user tiers",
          "AI services metered separately for fair-use pricing",
          "Single-tenant deployments quoted on infrastructure footprint",
          "Implementation services billed per phase, fixed-fee"]),
        ("Slide 12 — Why Glimmora",
         ["Purpose-built for GRC — not a generic workflow tool",
          "Founders bring decades of audit and consulting experience",
          "Modern stack (Next.js 16, Python AI) — not a 15-year-old legacy product",
          "Reference customers across BFSI, healthcare, and public sector"]),
        ("Slide 13 — Roadmap (12 Months)",
         ["Continuous-control monitoring integrations",
          "Native mobile auditee app",
          "ServiceNow / Jira CAPA webhooks",
          "Advanced analytics over multi-year audit history",
          "Voice-driven evidence capture for field auditors"]),
        ("Slide 14 — Call to Action",
         ["30-minute discovery call",
          "60-minute live demo on your audit universe",
          "Two-week pilot on UAT with a real engagement",
          "Production go-live in 6–8 weeks"]),
        ("Slide 15 — Q & A",
         ["Open the floor to questions",
          "Park advanced technical questions for a follow-up technical session"]),
    ]

    for title, bullets in slides:
        add_section_heading(doc, title, 1)
        add_bullets(doc, bullets)
        p = doc.add_paragraph()
        add_horizontal_rule(p)


def build_ia_usp(doc: Document):
    add_cover(doc, "Internal Audit Platform", "Unique Selling Propositions",
              "What makes this platform different — and why that matters")

    add_section_heading(doc, "1. Executive Summary", 1)
    add_paragraph(doc,
        "Internal audit tools tend to be one of two things: generic workflow systems "
        "bent into shape, or aging monoliths that predate modern web architecture. "
        "Our platform sits between those two extremes — purpose-built for internal "
        "audit, but on a modern stack that supports AI, multi-tenancy, and Arabic / "
        "Latvian RTL/LTR from day one. This document outlines the specific points "
        "of difference and the business value each one creates.")

    add_section_heading(doc, "2. USP Summary Table", 1)
    add_table(doc,
        ["#", "USP", "Buyer Value"],
        [
            ["1", "Risk-driven audit universe", "Defensible, board-ready coverage story"],
            ["2", "AI-augmented audit lifecycle", "Time saved at every stage without losing human control"],
            ["3", "Native multi-tenancy with single-tenant option", "Same product, your choice of deployment"],
            ["4", "End-to-end CAPA management with AI closure review", "Findings actually close — and stay closed"],
            ["5", "Document Library with full-text + AI search", "Past audits become a usable knowledge base"],
            ["6", "Multilingual UI + dynamic data translation", "True regional deployment, not a translation project"],
            ["7", "Encryption + RBAC built in from day one", "Pass infosec review without bolt-ons"],
            ["8", "Self-service auditee dashboard", "Cuts evidence chase time dramatically"],
            ["9", "Modern stack — Next.js 16 + Python AI", "Future-proof, hireable, fast to extend"],
            ["10", "Single-branch deploy with UAT/Prod gating", "Faster release cycles with safer rollouts"],
        ],
        col_widths=[1.0, 5.0, 9.5])

    add_section_heading(doc, "3. USPs in Detail", 1)

    usp_details = [
        ("3.1 Risk-Driven Audit Universe",
         "Most legacy tools store risks and audits as parallel lists. Our platform "
         "binds the Risk Universe and Audit Universe so that every engagement on the "
         "annual plan can be traced back to the risks it is intended to address. "
         "AI-Recommended Audits surface gaps where high-rated risks have no recent "
         "coverage.",
         "Audit Heads can defend their plan to the Audit Committee with evidence, not narrative."),
        ("3.2 AI-Augmented Lifecycle",
         "AI is embedded in the workflow — not bolted on as a side panel. AI suggests "
         "risks from process descriptions, drafts workpapers based on the process and "
         "risk, answers questions over ingested evidence, and reviews CAPA closure "
         "evidence for sufficiency. Every AI action is reviewable and reversible.",
         "Productivity gains of 30–50% on workpaper drafting and evidence review, "
         "without sacrificing auditor judgement."),
        ("3.3 Multi-Tenant by Design",
         "The data model is multi-tenant from the schema up — every business entity "
         "carries a customerAccountId that is enforced at the query layer. The same "
         "product can be deployed as a single-tenant instance on the customer's cloud "
         "with no code fork.",
         "Buyers don't choose between shared SaaS and bespoke deployment — they get both."),
        ("3.4 CAPA Lifecycle That Closes",
         "Most platforms track CAPA as a status column. Ours treats CAPA as a first-class "
         "workflow with owner, due date, evidence attachments, and an AI review step "
         "that flags weak or off-topic closure evidence before the auditor accepts it.",
         "Repeat findings drop because closures are substantive, not cosmetic."),
        ("3.5 Document Library With AI Search",
         "Every workpaper, finding, report, and evidence file is indexed by the Python "
         "AI backend. Auditors can ask natural-language questions across years of past "
         "audits and get cited passages back.",
         "Onboarding new auditors is dramatically faster; corporate audit memory survives "
         "staff turnover."),
        ("3.6 Multilingual From Day One",
         "English, Arabic (RTL), and Latvian (LTR) ship in the box. Static UI uses a "
         "phrase-based translation system; user-entered data is translated dynamically "
         "by the AI backend on create/edit and cached. New languages can be added "
         "without code changes.",
         "Buyers in the Gulf, Baltics, and other regional markets get a native experience, "
         "not a half-translated US product."),
        ("3.7 Security Built In, Not Bolted On",
         "AES-256-GCM field-level encryption, TLS 1.2+, RBAC enforced at three layers, "
         "and a documented 90-day master-key rotation runbook. Encryption can be toggled "
         "per environment so UAT and Production can be validated independently.",
         "Passes most infosec reviews on first pass; reduces sales cycle by weeks."),
        ("3.8 Self-Service Auditee Dashboard",
         "Auditees see every open evidence request, finding, and CAPA assigned to them "
         "in a single view, with context (related control, test objective, acceptance "
         "criteria) attached to each request.",
         "Less chasing by the audit team, faster fieldwork, better relationships between "
         "audit and business."),
        ("3.9 Modern Engineering Stack",
         "Next.js 16 App Router, React 19, Prisma 6, PostgreSQL, Python FastAPI, and a "
         "MongoDB-backed AI backend. Tooling and talent are widely available, and the "
         "stack is engineered for fast iteration.",
         "Customisations and integrations cost less and ship faster than on legacy "
         "platforms; no lock-in to a niche language or framework."),
        ("3.10 UAT/Prod Gating From a Single Branch",
         "A single GRC-MultiTenant branch deploys to both UAT and Production apps. "
         "Risky features ship dark behind per-app environment flags, are validated in "
         "UAT for 24–48 hours, then enabled in Production.",
         "Faster release cadence with measurably lower change-failure rate."),
    ]

    for h, body, value in usp_details:
        add_section_heading(doc, h, 2)
        add_paragraph(doc, body)
        add_callout(doc, "Why it matters", value)

    add_section_heading(doc, "4. Competitive Positioning", 1)
    add_table(doc,
        ["Capability", "Generic Workflow Tools", "Legacy GRC Suites", "Glimmora IA Platform"],
        [
            ["Purpose-built IA model", "No", "Yes", "Yes"],
            ["Modern web stack", "Often", "Rarely", "Yes (Next.js 16)"],
            ["Embedded AI", "Bolt-on", "Bolt-on", "Native"],
            ["Multilingual (incl. RTL)", "Partial", "Partial", "Native, 3 languages"],
            ["Multi-tenant + single-tenant", "Single only", "Single only", "Both"],
            ["Encryption + RBAC out of the box", "Partial", "Yes", "Yes"],
            ["Single-branch UAT/Prod gating", "Manual", "Manual", "Yes"],
        ],
        col_widths=[5.5, 3.0, 3.0, 4.0])

    add_section_heading(doc, "5. Closing Statement", 1)
    add_paragraph(doc,
        "The Glimmora Internal Audit Platform is not the cheapest tool, nor is it the "
        "most established brand. It is, however, the platform that maps most "
        "completely to how a modern internal audit function actually operates — "
        "risk-driven, AI-augmented, multilingual, and secure by default. For "
        "organisations that take assurance seriously, that combination is rare and "
        "valuable.")


# ============================================================
# CONTENT — GRC (Organization + Compliance + Risk + Assets)
# ============================================================

def build_grc_target_audience(doc: Document):
    add_cover(doc, "GRC Platform", "Target Audience",
              "Governance, Risk, and Compliance — who we serve and the value we deliver")

    add_section_heading(doc, "1. Purpose of This Document", 1)
    add_paragraph(doc,
        "This document identifies the primary and secondary audiences for the Glimmora "
        "GRC Platform. The GRC platform covers Organization (profile, context, "
        "processes, BIA), Compliance (frameworks, controls, governance, evidence, "
        "exceptions, KPIs), Risk Management (register, assessment, response, RCM), "
        "and Asset Management (inventory, classification). It is intended for solution "
        "buyers, internal champions, sales engineers, and onboarding teams.")

    add_section_heading(doc, "2. Primary Audiences", 1)

    add_section_heading(doc, "2.1 Chief Risk Officer (CRO) / Head of Risk", 2)
    add_paragraph(doc,
        "Owns the enterprise risk framework, reports to the Risk Committee, and "
        "needs a single view of risk exposure across the organisation.")
    add_paragraph(doc, "Key pain points addressed:", bold=True)
    add_bullets(doc, [
        "Risk register lives in spreadsheets — stale within days",
        "No shared taxonomy between risk, compliance, and audit",
        "Risk response actions are tracked in email, not the register",
        "KPIs and risk indicators sit in BI tools disconnected from the risk record",
    ])
    add_paragraph(doc, "Platform value:", bold=True)
    add_bullets(doc, [
        "Unified risk register with assessment, response, and Risk-Control Matrix",
        "Shared taxonomy across Risk, Compliance, and Internal Audit",
        "KPIs linked directly to risks and controls",
        "Heat-maps and dashboards driven from live data, not exports",
    ])

    add_section_heading(doc, "2.2 Chief Compliance Officer (CCO) / Compliance Head", 2)
    add_paragraph(doc,
        "Maintains compliance with multiple regulatory and voluntary frameworks "
        "(ISO 27001, SOC 2, PCI, GDPR, sector-specific regulators).")
    add_paragraph(doc, "Key pain points addressed:", bold=True)
    add_bullets(doc, [
        "Same control mapped to five frameworks — and re-tested five times",
        "Evidence collection is a once-a-year scramble before external audit",
        "Governance documents have no review schedule — they go stale silently",
        "Exceptions and waivers live outside any controlled register",
    ])
    add_paragraph(doc, "Platform value:", bold=True)
    add_bullets(doc, [
        "Framework Library with controls mapped across multiple frameworks (test once, satisfy many)",
        "Continuous evidence collection with assignees and due-date reminders",
        "Governance document register with review cadence and notification on expiry",
        "Exception register with owner, expiry, and approver workflow",
        "Statement of Applicability (SOA) and Regulatory Intelligence profiles",
    ])

    add_section_heading(doc, "2.3 Risk Owner / Risk Champion (Business)", 2)
    add_paragraph(doc,
        "Owns risks within their function (Finance, IT, Operations, HR, etc.) and is "
        "accountable for assessment, treatment, and reporting on those risks.")
    add_paragraph(doc, "Key pain points addressed:", bold=True)
    add_bullets(doc, [
        "Risk assessment is an annual ritual, not an operational habit",
        "Treatment actions are not visible to the second line",
        "No simple way to escalate emerging risks to the CRO",
    ])
    add_paragraph(doc, "Platform value:", bold=True)
    add_bullets(doc, [
        "Risk record with embedded assessment, response, and review history",
        "Treatment actions assigned, tracked, and visible to the CRO",
        "Escalation rules trigger when residual risk exceeds tolerance",
    ])

    add_section_heading(doc, "2.4 Control Owner / Process Owner", 2)
    add_paragraph(doc,
        "Responsible for designing, operating, and evidencing controls that mitigate "
        "risks and satisfy compliance requirements.")
    add_paragraph(doc, "Key pain points addressed:", bold=True)
    add_bullets(doc, [
        "Asked for the same evidence by multiple framework auditors",
        "Unclear which controls they actually own",
        "No way to flag control design issues without escalating",
    ])
    add_paragraph(doc, "Platform value:", bold=True)
    add_bullets(doc, [
        "Per-owner control dashboard with operating effectiveness status",
        "Evidence library — submit once, satisfy many frameworks",
        "Built-in exception workflow when controls cannot operate as designed",
    ])

    add_section_heading(doc, "2.5 Business Continuity / BIA Owner", 2)
    add_paragraph(doc,
        "Maintains the Business Impact Analysis, RTO/RPO targets, and BCP documentation.")
    add_paragraph(doc, "Key pain points addressed:", bold=True)
    add_bullets(doc, [
        "BIA is a 200-page Word document with no review trigger",
        "Process criticality is decided ad-hoc, not from a methodology",
        "No link between BIA and the underlying processes and assets",
    ])
    add_paragraph(doc, "Platform value:", bold=True)
    add_bullets(doc, [
        "Structured BIA per process with configurable methodology and categories",
        "Process catalogue tied to BIA, KPIs, and risks",
        "BCP labels and review schedules driven by criticality",
    ])

    add_section_heading(doc, "2.6 IT Asset / Information Asset Manager", 2)
    add_paragraph(doc,
        "Maintains the inventory of information assets and their classification per "
        "the organisation's data classification policy.")
    add_paragraph(doc, "Key pain points addressed:", bold=True)
    add_bullets(doc, [
        "Asset inventory is a CMDB export — out of date, IT-only",
        "Classification is inconsistent across departments",
        "No link between asset and the risks / controls applied to it",
    ])
    add_paragraph(doc, "Platform value:", bold=True)
    add_bullets(doc, [
        "Asset Inventory and Classification with custom attributes",
        "My Inventory view for individual asset custodians",
        "Asset-to-risk and asset-to-control linking",
    ])

    add_section_heading(doc, "3. Secondary Audiences", 1)

    add_section_heading(doc, "3.1 Board / Risk Committee", 2)
    add_bullets(doc, [
        "Live risk dashboards instead of static quarterly PDFs",
        "Trend analysis on residual risk, control failures, and overdue actions",
        "Drill-through from board chart to underlying record",
    ])

    add_section_heading(doc, "3.2 External Auditors / Regulators", 2)
    add_bullets(doc, [
        "Read-only access to evidence and SOA",
        "Auditable trail of every change to a control, evidence item, or exception",
        "Export packs for offline review",
    ])

    add_section_heading(doc, "3.3 IT / Information Security", 2)
    add_bullets(doc, [
        "AES-256-GCM at rest, TLS 1.2+ in transit",
        "RBAC with department-scoped roles",
        "Activity logs for SIEM ingestion",
    ])

    add_section_heading(doc, "4. Target Industries", 1)
    add_table(doc,
        ["Industry", "Why GRC Matters", "Specific Drivers"],
        [
            ["Banking & Financial Services", "Multi-framework, multi-regulator", "RBI, SEBI, Basel, ISO 27001"],
            ["Insurance", "ORSA, solvency, conduct risk", "IRDAI guidance, conduct frameworks"],
            ["Healthcare", "Patient data, accreditation, safety", "HIPAA-like, NABH/JCI"],
            ["Manufacturing", "Operational risk, supplier compliance", "ISO 9001/14001/45001"],
            ["IT / ITeS", "Client contracts, security certifications", "ISO 27001, SOC 2, GDPR, client audits"],
            ["Energy & Utilities", "Operational risk, safety, regulatory", "Sector regulators, ISO 55001"],
            ["Public Sector", "Compliance with statutory frameworks", "CAG, CVC, sector codes"],
        ],
        col_widths=[4.5, 5.5, 6.0])

    add_section_heading(doc, "5. Buying Personas", 1)
    add_table(doc,
        ["Persona", "Role in Decision", "What They Want to Hear"],
        [
            ["CRO", "Economic buyer for risk", "Live risk view, shared taxonomy, board-ready output"],
            ["CCO", "Economic buyer for compliance", "Multi-framework mapping, evidence reuse, audit-ready"],
            ["CFO", "Budget approver", "TCO, reduction in duplicated effort, fewer audit findings"],
            ["CISO", "Security gatekeeper", "Encryption, RBAC, secure SDLC, hosting options"],
            ["Process / Control Owner", "Daily user", "Less duplicate work, clear ownership, simple UI"],
        ],
        col_widths=[4.5, 4.5, 7.0])

    add_section_heading(doc, "6. Out of Scope (Not Our Audience)", 1)
    add_bullets(doc, [
        "Single-user consultants seeking a personal toolkit",
        "Organisations without a defined risk or compliance function",
        "Teams looking only for a policy document repository — we do more than that",
    ])


def build_grc_technical(doc: Document):
    add_cover(doc, "GRC Platform", "Technical Documentation",
              "Architecture, modules, data model, and integration points")

    add_section_heading(doc, "1. Solution Overview", 1)
    add_paragraph(doc,
        "The Glimmora GRC Platform is a multi-tenant SaaS application that unifies "
        "Organization context, Compliance, Risk Management, and Asset Management on a "
        "single Next.js 16 / PostgreSQL stack, augmented by a Python AI backend for "
        "document understanding and translation.")

    add_section_heading(doc, "2. Architecture", 1)
    add_table(doc,
        ["Layer", "Technology", "Responsibility"],
        [
            ["Presentation", "Next.js 16 App Router, React 19, Tailwind, shadcn/ui", "Server components, RBAC navigation, multilingual UI"],
            ["API", "Next.js Route Handlers + NextAuth v5 (JWT)", "REST endpoints, withAuth permission wrapper"],
            ["Business Logic", "TypeScript modules under src/lib", "Framework mapping, scoring, escalation, notifications"],
            ["Persistence", "Prisma 6 + PostgreSQL", "Multi-tenant schema with customerAccountId everywhere"],
            ["AI Services", "Python FastAPI + MongoDB", "Document ingestion, embeddings, RAG, translation"],
            ["File Storage", "S3-compatible object storage", "Evidence, governance docs, attachments"],
            ["Email", "SMTP relay", "Reminders, escalations, evidence requests"],
        ],
        col_widths=[3.5, 5.0, 7.5])

    add_section_heading(doc, "3. Module Inventory", 1)

    add_section_heading(doc, "3.1 Organization", 2)
    add_bullets(doc, [
        "Organisation Profile — legal entity, ownership, geography",
        "Organisation Context — ISO 9001/27001 style context-of-the-organisation register",
        "Process Catalogue — process register with owner, KPI, BIA",
        "BIA (Business Impact Analysis) — per-process structured BIA with configurable methodology and categories",
        "BCP Labels — criticality labels driven by BIA outcome",
        "User Management — users, departments, roles",
    ])

    add_section_heading(doc, "3.2 Compliance", 2)
    add_bullets(doc, [
        "Framework Library — ISO 27001, SOC 2, PCI, GDPR, and custom frameworks",
        "Control Library — controls mapped across multiple frameworks (test once, satisfy many)",
        "Statement of Applicability (SOA)",
        "Governance Documents — policies/procedures/standards with version, owner, review date",
        "Evidence — controlled evidence register with assignees and due-date reminders",
        "Exceptions — formal exception register with approver, expiry, and review",
        "KPIs — control / risk / compliance KPIs with target and actual",
        "Regulatory Intelligence — profiles tracking regulatory obligations",
        "Domain & Master Data — taxonomy, governance templates, evidence types",
    ])

    add_section_heading(doc, "3.3 Risk Management", 2)
    add_bullets(doc, [
        "Risk Register with assessment, response, and review history",
        "Risk-Control Matrix linking risks to mitigating controls",
        "Configurable scoring — impact, probability, scoring ranges",
        "Risk Universe and per-process risk catalogue",
        "AI-assisted risk identification from process narratives",
    ])

    add_section_heading(doc, "3.4 Asset Management", 2)
    add_bullets(doc, [
        "Asset Inventory with custom attributes",
        "Asset Classification per the organisation's data classification policy",
        "My Inventory — custodian view",
        "Asset-to-risk and asset-to-control linking",
        "Asset reports for management review",
    ])

    add_section_heading(doc, "4. Authentication & Authorization", 1)
    add_paragraph(doc, "RBAC enforced at navigation, API, and UI layers. Key GRC roles:")
    add_table(doc,
        ["Role", "Scope", "Typical Use"],
        [
            ["GRCAdministrator", "System / customer-account level", "Onboarding, framework setup"],
            ["CustomerAdministrator", "Organisation-wide admin", "Configuration, user management"],
            ["Reviewer", "Cross-module read + comment", "Risk / Compliance reviewers"],
            ["Contributor", "Edit assigned items", "Risk owners, control owners"],
            ["DepartmentReviewer", "Department-scoped read", "Functional heads"],
            ["DepartmentContributor", "Department-scoped edit", "Departmental risk / compliance owners"],
        ],
        col_widths=[5.0, 5.0, 5.5])

    add_section_heading(doc, "5. Data Model (Selected Entities)", 1)
    add_bullets(doc, [
        "Organization, Department, Process, ProcessKpi",
        "Bia, BiaCategory, BiaMethodology, BcpLabel",
        "Framework, Control, ControlMapping, Soa",
        "GovernanceDocument, GovernanceTemplate",
        "Evidence, EvidenceType",
        "Exception, Kpi, RegulatoryProfile",
        "Risk, RiskFactor, ImpactScale, ProbabilityScale, ScoringRange, RiskControlMatrix",
        "Asset, AssetClassification, AssetAttribute",
    ])

    add_section_heading(doc, "6. AI Services", 1)
    add_table(doc,
        ["Capability", "Description"],
        [
            ["Document Ingestion", "OCR, chunking, embeddings stored in MongoDB"],
            ["RAG over Governance Docs", "Natural-language Q&A on policies/standards"],
            ["Risk Suggestions", "Suggest risks from process descriptions or governance text"],
            ["Translation", "Dynamic translation of user-entered data into Arabic/Latvian on create/edit"],
        ],
        col_widths=[4.5, 11.0])

    add_section_heading(doc, "7. Security", 1)
    add_bullets(doc, [
        "TLS 1.2+ everywhere",
        "AES-256-GCM field-level encryption for sensitive Bytes columns",
        "Encryption kill switch via ENCRYPTION_ENABLED env var per app",
        "Bcrypt password hashing; configurable password policy",
        "Activity log for every state-changing API",
        "Secrets stored encrypted on DigitalOcean — never in git",
    ])
    add_callout(doc, "Compliance Mapping",
        "Maps to ISO 27001 A.5, A.6, A.8, A.9, A.10, A.12, A.18; SOC 2 CC1–CC9; "
        "GDPR Art. 32; PCI DSS Req. 3, 4, 7, 10.")

    add_section_heading(doc, "8. Integration Points", 1)
    add_table(doc,
        ["Integration", "Direction", "Notes"],
        [
            ["SSO (SAML / OIDC)", "Inbound", "Optional via NextAuth custom provider"],
            ["SMTP", "Outbound", "Reminders, escalations"],
            ["Object Storage (S3-compatible)", "Bi-directional", "Evidence and document storage"],
            ["BI / Analytics", "Outbound", "Read-replica or scheduled export"],
            ["SIEM", "Outbound", "Activity log streaming"],
            ["ITSM (ServiceNow / Jira)", "Outbound", "Action / exception webhooks (roadmap)"],
        ],
        col_widths=[4.5, 3.5, 8.0])

    add_section_heading(doc, "9. Scheduled Jobs", 1)
    add_bullets(doc, [
        "Daily 08:00 UTC — /api/cron/due-reminders: evidence due, policy reviews due, CAPA due",
        "Configurable escalation per evidence or policy item",
    ])

    add_section_heading(doc, "10. Internationalisation", 1)
    add_paragraph(doc,
        "English, Arabic (RTL), and Latvian (LTR) ship in the box. Static strings via "
        "phrase-based t(); user data via dynamic translation on create/edit only.")

    add_section_heading(doc, "11. Non-Functional Targets", 1)
    add_table(doc,
        ["Aspect", "Target"],
        [
            ["Page load (P95)", "< 2.0 s"],
            ["API response (P95)", "< 600 ms read, < 1.5 s write"],
            ["Concurrent users per tenant", "500+"],
            ["Uptime SLA", "99.5% UAT, 99.9% Prod"],
            ["RPO / RTO", "RPO 24 h, RTO 4 h"],
        ],
        col_widths=[5.0, 10.5])

    add_section_heading(doc, "12. Glossary", 1)
    add_table(doc,
        ["Term", "Meaning"],
        [
            ["SOA", "Statement of Applicability"],
            ["BIA", "Business Impact Analysis"],
            ["BCP", "Business Continuity Plan"],
            ["RCM", "Risk-Control Matrix"],
            ["RBAC", "Role-Based Access Control"],
        ],
        col_widths=[3.5, 12.0])


def build_grc_ppt(doc: Document):
    add_cover(doc, "GRC Platform", "Presentation Deck",
              "Slide-by-slide narrative for stakeholder briefings")

    slides = [
        ("Slide 1 — The Problem",
         ["Risk lives in spreadsheets, compliance lives in SharePoint, assets live in CMDB",
          "Same control is tested five times for five frameworks",
          "BIA is a 200-page doc nobody reads until DR happens",
          "Board asks 'What is our top risk?' and gets five different answers"]),
        ("Slide 2 — Why Now",
         ["Regulators require integrated risk + compliance + control evidence",
          "Customers demand SOC 2 / ISO 27001 in their procurement RFPs",
          "AI now usable for policy Q&A, risk suggestion, and translation",
          "Hybrid workforces need a shared system of record, not shared drives"]),
        ("Slide 3 — Our Solution",
         ["One platform: Organization + Compliance + Risk + Assets",
          "Shared taxonomy across all four pillars",
          "AI-augmented and multilingual (English / Arabic / Latvian)",
          "Multi-tenant SaaS or single-tenant in customer cloud"]),
        ("Slide 4 — The Four Pillars",
         ["Organization — profile, context, process catalogue, BIA",
          "Compliance — frameworks, controls, governance, evidence, exceptions, KPIs",
          "Risk — register, assessment, response, RCM",
          "Assets — inventory, classification, custodian view, links to risk + controls"]),
        ("Slide 5 — Compliance: Test Once, Satisfy Many",
         ["Map a control to ISO 27001, SOC 2, PCI, and GDPR simultaneously",
          "Evidence collected once is credited to every framework",
          "Statement of Applicability stays in sync automatically",
          "External audit prep collapses from weeks to days"]),
        ("Slide 6 — Risk: Live Register, Not a Spreadsheet",
         ["Risk record with embedded assessment, response, and history",
          "Risk-Control Matrix shows mitigation coverage",
          "Configurable impact / probability / scoring ranges",
          "AI suggests risks from process and governance text"]),
        ("Slide 7 — Organization & BIA",
         ["Process catalogue is the spine of the platform",
          "BIA per process with configurable methodology",
          "BCP labels drive criticality across risk + audit",
          "KPIs link to processes and controls"]),
        ("Slide 8 — Asset Management",
         ["Inventory with custom attributes and classification",
          "My Inventory view for custodians",
          "Linked to risks and controls — full traceability",
          "Reports for management review"]),
        ("Slide 9 — AI Capabilities",
         ["RAG over governance documents — chat with your policies",
          "Risk suggestions from process narratives",
          "Document ingestion with OCR + embeddings",
          "Dynamic translation of user-entered data"]),
        ("Slide 10 — Security & Trust",
         ["AES-256-GCM at rest, TLS 1.2+ in transit",
          "RBAC with department-scoped roles",
          "Activity log for every change",
          "Maps to ISO 27001, SOC 2, GDPR, PCI"]),
        ("Slide 11 — Deployment Options",
         ["Multi-tenant SaaS on DigitalOcean (default)",
          "Single-tenant in customer cloud (AWS / Azure / GCP)",
          "Hybrid — app in customer cloud, AI services on Glimmora backbone",
          "UAT + Production environments included with feature-flag gating"]),
        ("Slide 12 — Stakeholder Value",
         ["CRO — live risk dashboard, shared taxonomy",
          "CCO — multi-framework mapping, evidence reuse",
          "Control / Process Owner — clear ownership, less duplicate work",
          "CFO — fewer audit findings, lower TCO than separate point tools"]),
        ("Slide 13 — Implementation Approach",
         ["Week 1–2: Discovery and process / framework scoping",
          "Week 3–4: Configuration — frameworks, taxonomy, roles, departments",
          "Week 5–6: Pilot with one function (e.g., IT or Finance)",
          "Week 7–8: Migration of historical risks, controls, evidence",
          "Week 9+: Full rollout + 90 days of weekly office hours"]),
        ("Slide 14 — Pricing Model",
         ["Per-tenant annual subscription tied to user tiers",
          "AI services metered separately for fair-use pricing",
          "Single-tenant deployments quoted on infrastructure footprint",
          "Implementation services billed per phase, fixed-fee"]),
        ("Slide 15 — Why Glimmora",
         ["GRC purpose-built — not a generic workflow tool",
          "Founders bring decades of GRC and consulting experience",
          "Modern stack (Next.js 16, Python AI)",
          "Native multilingual including Arabic RTL"]),
        ("Slide 16 — Roadmap (12 Months)",
         ["Continuous-control monitoring connectors",
          "Native mobile evidence capture",
          "ITSM webhooks (ServiceNow / Jira)",
          "Quantitative risk (Monte Carlo) module",
          "Vendor-risk deep integration with TPRM platform"]),
        ("Slide 17 — Call to Action",
         ["30-minute discovery call",
          "60-minute live demo on your frameworks and risks",
          "Two-week pilot on UAT with a real function",
          "Production go-live in 6–8 weeks"]),
        ("Slide 18 — Q & A",
         ["Open the floor",
          "Park advanced technical questions for a follow-up session"]),
    ]

    for title, bullets in slides:
        add_section_heading(doc, title, 1)
        add_bullets(doc, bullets)
        p = doc.add_paragraph()
        add_horizontal_rule(p)


def build_grc_usp(doc: Document):
    add_cover(doc, "GRC Platform", "Unique Selling Propositions",
              "What makes this GRC platform different — and why that matters")

    add_section_heading(doc, "1. Executive Summary", 1)
    add_paragraph(doc,
        "Most GRC tools fall into one of three traps: they cover one pillar well "
        "(usually compliance) and bolt on the others; they are configurable to the "
        "point of being unusable; or they are aged monoliths that predate modern web "
        "engineering. Our platform avoids all three by covering Organization, "
        "Compliance, Risk, and Assets natively on a single modern stack, with shared "
        "taxonomy and embedded AI.")

    add_section_heading(doc, "2. USP Summary Table", 1)
    add_table(doc,
        ["#", "USP", "Buyer Value"],
        [
            ["1", "Four pillars on one stack with shared taxonomy", "No data silos, no duplicated taxonomies"],
            ["2", "Test once, satisfy many (control-framework mapping)", "Massive compliance leverage"],
            ["3", "Risk-Control Matrix built into the data model", "Coverage and gaps are visible by design"],
            ["4", "Structured BIA tied to process catalogue", "Continuity is data-driven, not a Word doc"],
            ["5", "Asset inventory linked to risks and controls", "Full chain of accountability"],
            ["6", "AI for policy Q&A, risk suggestion, translation", "Real productivity, not theatre"],
            ["7", "Native multilingual (EN / AR-RTL / LV)", "Regional deployment without translation projects"],
            ["8", "Multi-tenant + single-tenant from same codebase", "Same product, your choice of deployment"],
            ["9", "Encryption + RBAC built in", "Passes infosec review on first pass"],
            ["10", "Modern stack with single-branch UAT/Prod gating", "Faster cadence, safer rollouts"],
        ],
        col_widths=[1.0, 5.5, 9.0])

    add_section_heading(doc, "3. USPs in Detail", 1)

    usp_details = [
        ("3.1 Four Pillars, One Stack, One Taxonomy",
         "Organization, Compliance, Risk, and Assets share the same process, "
         "department, owner, and asset references. A change in one pillar propagates "
         "to the others automatically.",
         "No more 'what is the latest version of the risk register' debates."),
        ("3.2 Test Once, Satisfy Many",
         "Controls are mapped across frameworks in the Control Library. Evidence "
         "collected for one framework is automatically credited against every other "
         "framework that maps to the same control.",
         "External-audit prep collapses from weeks to days; duplicated evidence "
         "collection is eliminated."),
        ("3.3 Risk-Control Matrix Built In",
         "The RCM is a first-class table, not a report. Every risk shows its "
         "mitigating controls and every control shows the risks it covers. Gaps are "
         "queryable and visible on dashboards.",
         "CROs can prove coverage to the Board, not just claim it."),
        ("3.4 BIA Tied to the Process Catalogue",
         "Business Impact Analysis is captured per process using a configurable "
         "methodology and categories. BCP labels are derived from BIA outcomes and "
         "propagate to risk and audit prioritisation.",
         "Continuity planning becomes operational, not a once-a-year exercise."),
        ("3.5 Asset Inventory With Real Linkages",
         "Assets are not a parallel inventory — they are linked to the processes "
         "they support, the risks that threaten them, and the controls that protect "
         "them. My Inventory gives custodians a personal view.",
         "Information-security audits move faster because the chain is intact."),
        ("3.6 AI That Pulls Its Weight",
         "RAG over governance documents lets staff ask 'What is our incident "
         "response SLA?' and get a cited answer. Risk suggestions surface risks the "
         "team missed. Translation is automatic on create/edit.",
         "AI saves hours per user per week, every week — not just at demo time."),
        ("3.7 Multilingual From Day One",
         "English, Arabic (RTL), and Latvian (LTR) ship in the box. New languages "
         "can be added without code changes.",
         "Native experience for Gulf and Baltic markets — a competitive moat in "
         "those regions."),
        ("3.8 Multi-Tenant + Single-Tenant",
         "The data model is multi-tenant from the schema up. The same product can "
         "be deployed as single-tenant on the customer's cloud with no code fork.",
         "Buyers don't have to choose between shared SaaS and bespoke hosting."),
        ("3.9 Security Built In",
         "AES-256-GCM field-level encryption, TLS 1.2+, RBAC at three layers, and "
         "a documented 90-day key rotation runbook. Encryption can be toggled per "
         "environment.",
         "Passes most infosec reviews on first pass; reduces sales cycle by weeks."),
        ("3.10 Modern Stack + Single-Branch Gating",
         "Next.js 16, React 19, Prisma 6, PostgreSQL, Python FastAPI. A single "
         "GRC-MultiTenant branch deploys to both UAT and Production with per-app "
         "feature flags.",
         "Faster release cadence with lower change-failure rate; cheaper "
         "customisations than legacy platforms."),
    ]

    for h, body, value in usp_details:
        add_section_heading(doc, h, 2)
        add_paragraph(doc, body)
        add_callout(doc, "Why it matters", value)

    add_section_heading(doc, "4. Competitive Positioning", 1)
    add_table(doc,
        ["Capability", "Compliance-only Tools", "Legacy GRC Suites", "Glimmora GRC Platform"],
        [
            ["All four pillars natively", "No", "Partial", "Yes"],
            ["Shared taxonomy", "No", "Partial", "Yes"],
            ["Test-once-satisfy-many control mapping", "Partial", "Yes", "Yes"],
            ["BIA tied to process catalogue", "No", "Partial", "Yes"],
            ["Embedded AI", "Bolt-on", "Bolt-on", "Native"],
            ["Multilingual (incl. RTL)", "Partial", "Partial", "Native, 3 languages"],
            ["Multi-tenant + single-tenant", "Single only", "Single only", "Both"],
            ["Modern web stack", "Sometimes", "Rarely", "Yes (Next.js 16)"],
        ],
        col_widths=[5.5, 3.0, 3.0, 4.0])

    add_section_heading(doc, "5. Closing Statement", 1)
    add_paragraph(doc,
        "GRC is not a feature you bolt on — it is a discipline. The Glimmora GRC "
        "Platform respects that by offering one shared system for the four pillars "
        "that actually constitute GRC, built on a stack and a security posture that "
        "your infosec team will sign off on. For organisations that take governance "
        "seriously, that combination is rare and valuable.")


# ============================================================
# Main
# ============================================================

def build_doc(builder, out_path: Path):
    doc = Document()
    configure_styles(doc)

    # Margins
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)

    builder(doc)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    print(f"  wrote {out_path.relative_to(out_path.parent.parent.parent)}  ({out_path.stat().st_size // 1024} KB)")


def main():
    root = Path(__file__).resolve().parent.parent
    out_dir = root / "docs" / "platform-docs"
    print(f"Output directory: {out_dir}")

    jobs = [
        # Internal Audit (presentation is generated as .pptx — see generate-platform-presentations.py)
        (build_ia_target_audience,  out_dir / "InternalAudit-Target-Audience.docx"),
        (build_ia_technical,        out_dir / "InternalAudit-Technical-Documentation.docx"),
        (build_ia_usp,              out_dir / "InternalAudit-USP.docx"),
        # GRC (presentation is generated as .pptx — see generate-platform-presentations.py)
        (build_grc_target_audience, out_dir / "GRC-Target-Audience.docx"),
        (build_grc_technical,       out_dir / "GRC-Technical-Documentation.docx"),
        (build_grc_usp,             out_dir / "GRC-USP.docx"),
    ]

    for builder, path in jobs:
        build_doc(builder, path)

    print(f"\nDone. {len(jobs)} documents generated in {out_dir}")


if __name__ == "__main__":
    main()
